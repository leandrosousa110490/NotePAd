import sys
import os
from PyQt6.QtWidgets import (QApplication, QMainWindow, QTextEdit, QFileDialog, 
                             QMessageBox, QFontDialog, QColorDialog, QVBoxLayout, 
                             QWidget, QMenuBar, QMenu, QToolBar, QStatusBar,
                             QInputDialog, QDialog, QDialogButtonBox, QSpinBox, 
                             QLabel, QHBoxLayout, QCheckBox, QLineEdit, QPushButton,
                             QTabWidget, QPlainTextEdit, QComboBox, QSlider, QSplitter,
                             QScrollArea)
from PyQt6.QtGui import (QAction, QFont, QTextCursor, QTextCharFormat, 
                         QColor, QKeySequence, QTextTableFormat, QTextImageFormat,
                         QTextListFormat, QImage, QTextFrameFormat, QSyntaxHighlighter,
                         QTextDocument, QPalette, QTextBlockFormat, QTextLength,
                         QTextTableCellFormat, QPainter, QTextFormat, QPixmap)
from PyQt6.QtCore import Qt, QSettings, QRegularExpression, QTimer, QDateTime, QRect, QSize, QProcess, QEvent
from PyQt6.QtPrintSupport import QPrintDialog, QPrinter
import re
import subprocess
import tempfile
import json
import shutil
import uuid
try:
    import duckdb
except Exception:
    duckdb = None
try:
    import polars as pl
except Exception:
    pl = None
try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None


class PythonHighlighter(QSyntaxHighlighter):
    def __init__(self, document):
        super().__init__(document)
        self.highlighting_rules = []
        
        # Keywords
        keyword_format = QTextCharFormat()
        keyword_format.setForeground(QColor("#CC7832"))
        keyword_format.setFontWeight(QFont.Weight.Bold)
        keywords = [
            "and", "as", "assert", "break", "class", "continue", "def",
            "del", "elif", "else", "except", "False", "finally", "for",
            "from", "global", "if", "import", "in", "is", "lambda", "None",
            "nonlocal", "not", "or", "pass", "raise", "return", "True",
            "try", "while", "with", "yield"
        ]
        for word in keywords:
            pattern = QRegularExpression(f"\\b{word}\\b")
            self.highlighting_rules.append((pattern, keyword_format))
        
        # Strings
        string_format = QTextCharFormat()
        string_format.setForeground(QColor("#6A8759"))
        self.highlighting_rules.append((QRegularExpression('"[^"]*"'), string_format))
        self.highlighting_rules.append((QRegularExpression("'[^']*'"), string_format))
        
        # Comments
        comment_format = QTextCharFormat()
        comment_format.setForeground(QColor("#808080"))
        comment_format.setFontItalic(True)
        self.highlighting_rules.append((QRegularExpression("#[^\n]*"), comment_format))
        
    def highlightBlock(self, text):
        for pattern, format in self.highlighting_rules:
            iterator = pattern.globalMatch(text)
            while iterator.hasNext():
                match = iterator.next()
                self.setFormat(match.capturedStart(), match.capturedLength(), format)


class TableDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Insert Table")
        self.setModal(True)
        
        layout = QVBoxLayout()
        
        row_layout = QHBoxLayout()
        row_layout.addWidget(QLabel("Rows:"))
        self.rows_spin = QSpinBox()
        self.rows_spin.setMinimum(1)
        self.rows_spin.setMaximum(50)
        self.rows_spin.setValue(3)
        row_layout.addWidget(self.rows_spin)
        layout.addLayout(row_layout)
        
        col_layout = QHBoxLayout()
        col_layout.addWidget(QLabel("Columns:"))
        self.cols_spin = QSpinBox()
        self.cols_spin.setMinimum(1)
        self.cols_spin.setMaximum(20)
        self.cols_spin.setValue(3)
        col_layout.addWidget(self.cols_spin)
        layout.addLayout(col_layout)
        
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | 
                                   QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        
        self.setLayout(layout)


class FindReplaceDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_editor = parent
        self.setWindowTitle("Find and Replace")
        self.setModal(False)
        
        layout = QVBoxLayout()
        
        # Find input
        find_layout = QHBoxLayout()
        find_layout.addWidget(QLabel("Find:"))
        self.find_input = QLineEdit()
        find_layout.addWidget(self.find_input)
        layout.addLayout(find_layout)
        
        # Replace input
        replace_layout = QHBoxLayout()
        replace_layout.addWidget(QLabel("Replace:"))
        self.replace_input = QLineEdit()
        replace_layout.addWidget(self.replace_input)
        layout.addLayout(replace_layout)
        
        # Options
        self.case_sensitive = QCheckBox("Case sensitive")
        self.whole_word = QCheckBox("Whole words only")
        layout.addWidget(self.case_sensitive)
        layout.addWidget(self.whole_word)
        
        # Buttons
        button_layout = QHBoxLayout()
        find_btn = QPushButton("Find Next")
        find_btn.clicked.connect(self.find_next)
        replace_btn = QPushButton("Replace")
        replace_btn.clicked.connect(self.replace_current)
        replace_all_btn = QPushButton("Replace All")
        replace_all_btn.clicked.connect(self.replace_all)
        
        button_layout.addWidget(find_btn)
        button_layout.addWidget(replace_btn)
        button_layout.addWidget(replace_all_btn)
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        
    def find_next(self):
        if self.parent_editor and hasattr(self.parent_editor, 'current_tab'):
            text_edit = self.parent_editor.current_tab()
            if text_edit:
                search_text = self.find_input.text()
                flags = QTextDocument.FindFlag(0)
                if self.case_sensitive.isChecked():
                    flags |= QTextDocument.FindFlag.FindCaseSensitively
                if self.whole_word.isChecked():
                    flags |= QTextDocument.FindFlag.FindWholeWords
                
                if not text_edit.find(search_text, flags):
                    cursor = text_edit.textCursor()
                    cursor.movePosition(QTextCursor.MoveOperation.Start)
                    text_edit.setTextCursor(cursor)
                    if not text_edit.find(search_text, flags):
                        QMessageBox.information(self, "Find", "Text not found")
    
    def replace_current(self):
        if self.parent_editor and hasattr(self.parent_editor, 'current_tab'):
            text_edit = self.parent_editor.current_tab()
            if text_edit:
                cursor = text_edit.textCursor()
                if cursor.hasSelection():
                    cursor.insertText(self.replace_input.text())
                self.find_next()
    
    def replace_all(self):
        if self.parent_editor and hasattr(self.parent_editor, 'current_tab'):
            text_edit = self.parent_editor.current_tab()
            if text_edit:
                search_text = self.find_input.text()
                replace_text = self.replace_input.text()
                
                cursor = text_edit.textCursor()
                cursor.beginEditBlock()
                cursor.movePosition(QTextCursor.MoveOperation.Start)
                text_edit.setTextCursor(cursor)
                
                count = 0
                flags = QTextDocument.FindFlag(0)
                if self.case_sensitive.isChecked():
                    flags |= QTextDocument.FindFlag.FindCaseSensitively
                if self.whole_word.isChecked():
                    flags |= QTextDocument.FindFlag.FindWholeWords
                
                while text_edit.find(search_text, flags):
                    cursor = text_edit.textCursor()
                    cursor.insertText(replace_text)
                    count += 1
                
                cursor.endEditBlock()
                QMessageBox.information(self, "Replace All", f"Replaced {count} occurrence(s)")

class ImageResizeDialog(QDialog):
    def __init__(self, width, height, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Resize Image")
        self.setModal(True)
        layout = QVBoxLayout()
        w_layout = QHBoxLayout()
        w_layout.addWidget(QLabel("Width:"))
        self.width_spin = QSpinBox()
        self.width_spin.setMinimum(10)
        self.width_spin.setMaximum(5000)
        self.width_spin.setValue(int(width) if width else 100)
        w_layout.addWidget(self.width_spin)
        layout.addLayout(w_layout)
        h_layout = QHBoxLayout()
        h_layout.addWidget(QLabel("Height:"))
        self.height_spin = QSpinBox()
        self.height_spin.setMinimum(10)
        self.height_spin.setMaximum(5000)
        self.height_spin.setValue(int(height) if height else 100)
        h_layout.addWidget(self.height_spin)
        layout.addLayout(h_layout)
        self.lock_ratio = QCheckBox("Lock aspect ratio")
        self.lock_ratio.setChecked(True)
        layout.addWidget(self.lock_ratio)
        self._ratio = (width / height) if width and height else None
        def sync_height(val):
            if self.lock_ratio.isChecked() and self._ratio:
                self.height_spin.blockSignals(True)
                self.height_spin.setValue(max(10, int(round(val / self._ratio))))
                self.height_spin.blockSignals(False)
        def sync_width(val):
            if self.lock_ratio.isChecked() and self._ratio:
                self.width_spin.blockSignals(True)
                self.width_spin.setValue(max(10, int(round(val * self._ratio))))
                self.width_spin.blockSignals(False)
        self.width_spin.valueChanged.connect(sync_height)
        self.height_spin.valueChanged.connect(sync_width)
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        self.setLayout(layout)

class LanguageHighlighter(QSyntaxHighlighter):
    def __init__(self, document, language="Python"):
        super().__init__(document)
        self.language = language
        self.highlighting_rules = []
        self.multi_start = None
        self.multi_end = None
        self.set_language(language)
    def set_language(self, language):
        self.language = language
        self.highlighting_rules = []
        self.multi_start = None
        self.multi_end = None
        if language == "Text":
            return
        kw_map = {
            "Python": [
                "and","as","assert","break","class","continue","def","del","elif","else","except","False","finally","for","from","global","if","import","in","is","lambda","None","nonlocal","not","or","pass","raise","return","True","try","while","with","yield"
            ],
            "R": [
                "function","if","else","repeat","while","for","in","next","break","TRUE","FALSE","NULL","NA","NaN","Inf"
            ]
        }
        keyword_format = QTextCharFormat()
        keyword_format.setForeground(QColor("#CC7832"))
        keyword_format.setFontWeight(QFont.Weight.Bold)
        for w in kw_map.get(language, []):
            self.highlighting_rules.append((QRegularExpression(f"\\b{w}\\b"), keyword_format))
        string_format = QTextCharFormat()
        string_format.setForeground(QColor("#6A8759"))
        if language != "Text":
            self.highlighting_rules.append((QRegularExpression('"[^"]*"'), string_format))
            self.highlighting_rules.append((QRegularExpression("'[^']*'"), string_format))
        comment_format = QTextCharFormat()
        comment_format.setForeground(QColor("#808080"))
        comment_format.setFontItalic(True)
        if language in ("Python","R"):
            self.highlighting_rules.append((QRegularExpression("#[^\n]*"), comment_format))
    def highlightBlock(self, text):
        for pattern, fmt in self.highlighting_rules:
            it = pattern.globalMatch(text)
            while it.hasNext():
                m = it.next()
                self.setFormat(m.capturedStart(), m.capturedLength(), fmt)
        if self.multi_start and self.multi_end:
            start = 0
            if self.previousBlockState() != 1:
                match = self.multi_start.match(text, 0)
                start = match.capturedStart() if match.hasMatch() else -1
            else:
                start = 0
            while start >= 0:
                match_end = self.multi_end.match(text, start)
                end = match_end.capturedEnd() if match_end.hasMatch() else -1
                length = (end - start) if end >= 0 else (len(text) - start)
                fmt = QTextCharFormat()
                fmt.setForeground(QColor("#808080"))
                self.setFormat(start, length, fmt)
                if end < 0:
                    self.setCurrentBlockState(1)
                    return
                start_match = self.multi_start.match(text, end)
                start = start_match.capturedStart() if start_match.hasMatch() else -1
            self.setCurrentBlockState(0)

try:
    from pygments import lex
    from pygments.lexers import PythonLexer
    try:
        from pygments.lexers.r import SLexer as RLexer
    except Exception:
        RLexer = None
    from pygments.token import Token
except Exception:
    lex = None
    PythonLexer = None
    RLexer = None
    Token = None

class PygmentsHighlighter(QSyntaxHighlighter):
    def __init__(self, document, lexer_name="Python"):
        super().__init__(document)
        self.lexer = None
        self.language = lexer_name
        self.set_lexer_from_language(lexer_name)
    def set_lexer_from_language(self, language):
        self.language = language
        if language == "Python" and PythonLexer:
            self.lexer = PythonLexer()
        elif language == "R" and RLexer:
            self.lexer = RLexer()
        else:
            self.lexer = None
    def highlightBlock(self, text):
        if not self.lexer or not lex:
            return
        for tok_type, tok_text in lex(text, self.lexer):
            fmt = QTextCharFormat()
            if tok_type in Token.Keyword:
                fmt.setForeground(QColor("#CC7832"))
                fmt.setFontWeight(QFont.Weight.Bold)
            elif tok_type in Token.String:
                fmt.setForeground(QColor("#6A8759"))
            elif tok_type in Token.Comment:
                fmt.setForeground(QColor("#808080"))
                fmt.setFontItalic(True)
            elif tok_type in Token.Number:
                fmt.setForeground(QColor("#6897BB"))
            elif tok_type in Token.Name.Builtin:
                fmt.setForeground(QColor("#A5C261"))
            start = text.find(tok_text, 0 if tok_text else 0)
            while fmt.foreground().color().isValid() and tok_text and start != -1:
                self.setFormat(start, len(tok_text), fmt)
                start = text.find(tok_text, start + len(tok_text))

class LineNumberArea(QWidget):
    def __init__(self, editor):
        super().__init__(editor)
        self.editor = editor
    def sizeHint(self):
        return QSize(self.editor.lineNumberAreaWidth(), 0)
    def paintEvent(self, event):
        self.editor.lineNumberAreaPaintEvent(event)

class RichTextEdit(QTextEdit):
    def __init__(self):
        super().__init__()
        self.showLineNumbers = False
        self.lineNumberArea = LineNumberArea(self)
        self.verticalScrollBar().valueChanged.connect(self.lineNumberArea.update)
        self.textChanged.connect(self.lineNumberArea.update)
        self.cursorPositionChanged.connect(self.highlightCurrentLine)
        self.highlighter = PygmentsHighlighter(self.document(), "Python") if lex and PythonLexer else LanguageHighlighter(self.document(), "Python")
        self.updateLineNumberAreaWidth()
        self._img_resize_active = False
        self._img_cursor_pos = None
        self._img_start_pos = None
        self._img_initial_size = None
        self._img_name = None
        self._img_handle_pad = 8
    def canInsertFromMimeData(self, source):
        if hasattr(source, 'hasImage') and source.hasImage():
            return True
        if hasattr(source, 'hasUrls') and source.hasUrls():
            return True
        return super().canInsertFromMimeData(source)
    def insertFromMimeData(self, source):
        try:
            if hasattr(source, 'hasImage') and source.hasImage():
                image = source.imageData()
                if isinstance(image, QImage):
                    img = image
                else:
                    img = QImage(image)
                if not img.isNull():
                    max_width = 600
                    if img.width() > max_width:
                        img = img.scaledToWidth(max_width, Qt.TransformationMode.SmoothTransformation)
                    base_dir = getattr(self, 'session_images_dir', None) or tempfile.gettempdir()
                    os.makedirs(base_dir, exist_ok=True)
                    temp_path = os.path.join(base_dir, f"img_{uuid.uuid4().hex}.png")
                    img.save(temp_path, 'PNG')
                    fmt = QTextImageFormat()
                    fmt.setName(temp_path)
                    fmt.setWidth(img.width())
                    fmt.setHeight(img.height())
                    self.textCursor().insertImage(fmt)
                    return
            if hasattr(source, 'hasUrls') and source.hasUrls():
                from PyQt6.QtCore import QUrl
                urls = source.urls()
                for u in urls:
                    p = u.toLocalFile()
                    if p and os.path.exists(p):
                        ext = os.path.splitext(p)[1].lower()
                        if ext in ('.png','.jpg','.jpeg','.bmp','.gif'):
                            img = QImage(p)
                            if not img.isNull():
                                max_width = 600
                                if img.width() > max_width:
                                    img = img.scaledToWidth(max_width, Qt.TransformationMode.SmoothTransformation)
                                base_dir = getattr(self, 'session_images_dir', None) or tempfile.gettempdir()
                                os.makedirs(base_dir, exist_ok=True)
                                dest = os.path.join(base_dir, f"img_{uuid.uuid4().hex}{ext}")
                                img.save(dest)
                                fmt = QTextImageFormat()
                                fmt.setName(dest)
                                fmt.setWidth(img.width())
                                fmt.setHeight(img.height())
                                self.textCursor().insertImage(fmt)
                                continue
        except Exception:
            pass
        super().insertFromMimeData(source)
    def lineNumberAreaWidth(self):
        if not self.showLineNumbers:
            return 0
        digits = len(str(max(1, self.document().blockCount())))
        fm = self.fontMetrics()
        return 14 + fm.horizontalAdvance("9") * digits
    def updateLineNumberAreaWidth(self):
        self.setViewportMargins(self.lineNumberAreaWidth(), 0, 0, 0)
    def resizeEvent(self, event):
        super().resizeEvent(event)
        cr = self.contentsRect()
        self.lineNumberArea.setGeometry(QRect(cr.left(), cr.top(), self.lineNumberAreaWidth(), cr.height()))
    def lineNumberAreaPaintEvent(self, event):
        if not self.showLineNumbers:
            return
        painter = QPainter(self.lineNumberArea)
        painter.fillRect(event.rect(), self.palette().color(QPalette.ColorRole.Base))
        block = self.document().firstBlock()
        viewport_h = self.viewport().height()
        printed_tops = []
        line_no = 0
        while block.isValid():
            cursor = QTextCursor(block)
            r = self.cursorRect(cursor)
            y = r.top()
            already = any(abs(y - py) <= 1 for py in printed_tops)
            if not already:
                line_no += 1
            if y > viewport_h:
                break
            if r.bottom() >= 0 and not already:
                printed_tops.append(y)
                painter.setPen(QColor("#808080"))
                fm = self.fontMetrics()
                text = str(line_no)
                x = self.lineNumberArea.width() - fm.horizontalAdvance(text) - 6
                painter.drawText(x, y + fm.ascent(), text)
            block = block.next()
    def setLineNumbersVisible(self, visible):
        self.showLineNumbers = visible
        self.updateLineNumberAreaWidth()
        self.lineNumberArea.update()
        self.highlightCurrentLine()
    def highlightCurrentLine(self):
        if not self.showLineNumbers:
            self.setExtraSelections([])
            return
        sel = QTextEdit.ExtraSelection()
        line_color = self.palette().color(QPalette.ColorRole.Base).lighter(110)
        sel.format.setBackground(line_color)
        sel.format.setProperty(QTextFormat.Property.FullWidthSelection, True)
        sel.cursor = self.textCursor()
        sel.cursor.clearSelection()
        self.setExtraSelections([sel])
    def setLanguage(self, language):
        use_pygments = False
        if lex:
            if language == "Python" and PythonLexer:
                use_pygments = True
            elif language == "R" and RLexer:
                use_pygments = True
        if hasattr(self, 'highlighter') and self.highlighter:
            try:
                self.highlighter.setDocument(None)
            except Exception:
                pass
        if use_pygments:
            self.highlighter = PygmentsHighlighter(self.document(), language)
        else:
            self.highlighter = LanguageHighlighter(self.document(), language)
        try:
            self.highlighter.rehighlight()
        except Exception:
            pass
    def keyPressEvent(self, event):
        if (event.modifiers() & Qt.KeyboardModifier.AltModifier) and event.key() == Qt.Key.Key_Right:
            cursor = self.textCursor()
            table = cursor.currentTable()
            if table:
                cell = table.cellAt(cursor)
                if cell.isValid():
                    row = cell.row()
                    col = cell.column()
                    if col + 1 >= table.columns():
                        table.insertColumns(col + 1, 1)
                        text_color = self.palette().color(QPalette.ColorRole.Text)
                        cell_fmt = QTextTableCellFormat()
                        cell_fmt.setBorder(1)
                        cell_fmt.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Solid)
                        cell_fmt.setBorderBrush(text_color)
                        for r in range(table.rows()):
                            table.cellAt(r, col + 1).setFormat(cell_fmt)
                    target = table.cellAt(row, col + 1).firstCursorPosition()
                    self.setTextCursor(target)
                    event.accept()
                    return
        if (event.modifiers() & Qt.KeyboardModifier.AltModifier) and event.key() == Qt.Key.Key_Left:
            cursor = self.textCursor()
            table = cursor.currentTable()
            if table:
                cell = table.cellAt(cursor)
                if cell.isValid():
                    row = cell.row()
                    col = cell.column()
                    if col - 1 < 0:
                        table.insertColumns(0, 1)
                        text_color = self.palette().color(QPalette.ColorRole.Text)
                        cell_fmt = QTextTableCellFormat()
                        cell_fmt.setBorder(1)
                        cell_fmt.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Solid)
                        cell_fmt.setBorderBrush(text_color)
                        for r in range(table.rows()):
                            table.cellAt(r, 0).setFormat(cell_fmt)
                        target_col = 0
                    else:
                        target_col = col - 1
                    target = table.cellAt(row, target_col).firstCursorPosition()
                    self.setTextCursor(target)
                    event.accept()
                    return
        if (event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter)) and (event.modifiers() & Qt.KeyboardModifier.AltModifier):
            cursor = self.textCursor()
            table = cursor.currentTable()
            if table:
                cell = table.cellAt(cursor)
                if cell.isValid():
                    row = cell.row()
                    col = cell.column()
                    if row + 1 >= table.rows():
                        table.insertRows(row + 1, 1)
                        text_color = self.palette().color(QPalette.ColorRole.Text)
                        cell_fmt = QTextTableCellFormat()
                        cell_fmt.setBorder(1)
                        cell_fmt.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Solid)
                        cell_fmt.setBorderBrush(text_color)
                        for c in range(table.columns()):
                            table.cellAt(row + 1, c).setFormat(cell_fmt)
                    target = table.cellAt(row + 1, col).firstCursorPosition()
                    self.setTextCursor(target)
                    event.accept()
                    return
        super().keyPressEvent(event)
    def mousePressEvent(self, event):
        cursor = self.cursorForPosition(event.pos())
        fmt = cursor.charFormat()
        if fmt.isImageFormat() and event.button() == Qt.MouseButton.LeftButton:
            name = fmt.property(QTextFormat.Property.ImageName)
            w = fmt.property(QTextFormat.Property.ImageWidth)
            h = fmt.property(QTextFormat.Property.ImageHeight)
            if not w or not h:
                try:
                    img = QImage(name) if name else QImage()
                    w = int(img.width()) if img and img.width() > 0 else 100
                    h = int(img.height()) if img and img.height() > 0 else 100
                except Exception:
                    w, h = 100, 100
            r = self.cursorRect(cursor)
            ir = QRect(r.left(), r.top(), int(w), int(h))
            br = QRect(ir.right() - self._img_handle_pad, ir.bottom() - self._img_handle_pad, self._img_handle_pad*2, self._img_handle_pad*2)
            if br.contains(event.pos()):
                self._img_resize_active = True
                self._img_cursor_pos = cursor.position()
                self._img_start_pos = event.pos()
                self._img_initial_size = (int(w), int(h))
                self._img_name = name
                self.setCursor(Qt.CursorShape.SizeFDiagCursor)
                event.accept()
                return
        block = cursor.block()
        i = cursor.positionInBlock()
        t = block.text()
        if i < len(t):
            ch = t[i]
            if ch in ['☐', '☑']:
                c = QTextCursor(block)
                c.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.MoveAnchor, i)
                c.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor, 1)
                new_ch = '☑' if ch == '☐' else '☐'
                c.insertText(new_ch)
                return
        super().mousePressEvent(event)
    def mouseMoveEvent(self, event):
        if self._img_resize_active and self._img_cursor_pos is not None and self._img_initial_size is not None:
            dx = max(1, event.pos().x() - self._img_start_pos.x())
            dy = max(1, event.pos().y() - self._img_start_pos.y())
            new_w = max(10, self._img_initial_size[0] + dx)
            new_h = max(10, self._img_initial_size[1] + dy)
            c = self.textCursor()
            c.setPosition(self._img_cursor_pos)
            fmt = c.charFormat()
            if fmt.isImageFormat():
                fmt.setProperty(QTextFormat.Property.ImageWidth, int(new_w))
                fmt.setProperty(QTextFormat.Property.ImageHeight, int(new_h))
                c.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor, 1)
                c.mergeCharFormat(fmt)
                self.viewport().update()
            event.accept()
            return
        cursor = self.cursorForPosition(event.pos())
        fmt = cursor.charFormat()
        if fmt.isImageFormat():
            name = fmt.property(QTextFormat.Property.ImageName)
            w = fmt.property(QTextFormat.Property.ImageWidth) or 100
            h = fmt.property(QTextFormat.Property.ImageHeight) or 100
            r = self.cursorRect(cursor)
            ir = QRect(r.left(), r.top(), int(w), int(h))
            br = QRect(ir.right() - self._img_handle_pad, ir.bottom() - self._img_handle_pad, self._img_handle_pad*2, self._img_handle_pad*2)
            if br.contains(event.pos()):
                self.setCursor(Qt.CursorShape.SizeFDiagCursor)
            else:
                self.setCursor(Qt.CursorShape.IBeamCursor)
        else:
            self.setCursor(Qt.CursorShape.IBeamCursor)
        super().mouseMoveEvent(event)
    def mouseReleaseEvent(self, event):
        if self._img_resize_active:
            self._img_resize_active = False
            self._img_cursor_pos = None
            self._img_start_pos = None
            self._img_initial_size = None
            self._img_name = None
            self.setCursor(Qt.CursorShape.IBeamCursor)
        super().mouseReleaseEvent(event)


class AdvancedNotepad(QMainWindow):
    def __init__(self):
        super().__init__()
        self.current_file = None
        self.is_modified = False
        self.zoom_level = 100
        self.themes = {
            "Light": {"bg": "#FFFFFF", "fg": "#000000"},
            "Dark": {"bg": "#2B2B2B", "fg": "#A9B7C6"},
            "Monokai": {"bg": "#272822", "fg": "#F8F8F2"},
            "Solarized": {"bg": "#002B36", "fg": "#839496"}
        }
        self.python_temp_path = os.path.join(tempfile.gettempdir(), "advanced_notepad_temp.py")
        self.r_temp_path = os.path.join(tempfile.gettempdir(), "advanced_notepad_temp.R")
        self.session_dir = os.path.join(tempfile.gettempdir(), "advanced_notepad_session")
        self.images_dir = os.path.join(self.session_dir, "images")
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Ultra Advanced Notepad++")
        self.setGeometry(100, 100, 1200, 800)
        
        # Create tab widget for multiple documents
        self.tab_widget = QTabWidget()
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.close_tab)
        self.tab_widget.currentChanged.connect(self.tab_changed)
        self.setCentralWidget(self.tab_widget)
        self.tab_widget.tabBar().installEventFilter(self)
        
        # Create status bar before creating the first tab
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        
        self.restore_session_or_default()
        
        # Create menu bar
        self.create_menus()
        
        # Create toolbar
        self.create_toolbar()
        
        # Auto-save timer
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(60000)  # Auto-save every 60 seconds
        
        self.capture_active = False
        self.capture_timer = QTimer()
        self.capture_timer.setInterval(2000)
        self.capture_timer.timeout.connect(self._capture_tick)
        self.captured_images = []
        self._capture_filter_installed = False
        
        self.update_status()
        
    def new_tab(self):
        text_edit = RichTextEdit()
        text_edit.setFont(QFont("Consolas", 11))
        text_edit.textChanged.connect(self.text_changed)
        text_edit.cursorPositionChanged.connect(self.update_status)
        
        # Enable drag and drop for images
        text_edit.setAcceptDrops(True)
        if hasattr(text_edit, "setLanguage"):
            text_edit.setLanguage("Text")
        text_edit.session_images_dir = self.images_dir
        text_edit.setLineNumbersVisible(True)
        index = self.tab_widget.addTab(text_edit, "Untitled")
        self.tab_widget.setCurrentIndex(index)
        return text_edit
        
    def current_tab(self):
        w = self.tab_widget.currentWidget()
        if hasattr(w, 'text_editor'):
            return w.text_editor
        return w
        
    def close_tab(self, index):
        if self.tab_widget.count() > 1:
            self.tab_widget.removeTab(index)
        else:
            self.close()
            
    def tab_changed(self, index):
        self.update_status()
    def eventFilter(self, obj, event):
        if obj == self.tab_widget.tabBar() and event.type() == QEvent.Type.MouseButtonDblClick:
            idx = obj.tabAt(event.pos())
            if idx >= 0:
                current = self.tab_widget.tabText(idx)
                name, ok = QInputDialog.getText(self, "Rename Tab", "Tab name:", QLineEdit.EchoMode.Normal, current)
                if ok:
                    n = name.strip()
                    if n:
                        self.tab_widget.setTabText(idx, n)
            return True
        if getattr(self, 'capture_active', False) and event.type() == QEvent.Type.MouseButtonPress:
            try:
                self._capture_click()
            except Exception:
                pass
        return super().eventFilter(obj, event)
    def restore_session_or_default(self):
        if not self.restore_session():
            self.new_tab()
    def restore_session(self):
        try:
            meta_path = os.path.join(self.session_dir, "session.json")
            if not os.path.exists(meta_path):
                return False
            with open(meta_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            while self.tab_widget.count() > 0:
                self.tab_widget.removeTab(0)
            for item in data.get("tabs", []):
                editor = self.new_tab()
                idx = self.tab_widget.currentIndex()
                self.tab_widget.setTabText(idx, item.get("title", "Untitled"))
                lang = item.get("lang", "Text")
                if hasattr(editor, "setLanguage"):
                    editor.setLanguage(lang)
                file_path = os.path.join(self.session_dir, item.get("file", ""))
                if os.path.exists(file_path):
                    with open(file_path, "r", encoding="utf-8") as rf:
                        editor.setHtml(rf.read())
            return True
        except Exception:
            return False
    def save_session(self):
        try:
            os.makedirs(self.session_dir, exist_ok=True)
            os.makedirs(self.images_dir, exist_ok=True)
            tabs = []
            for i in range(self.tab_widget.count()):
                w = self.tab_widget.widget(i)
                editor = w.text_editor if hasattr(w, "text_editor") else w
                title = self.tab_widget.tabText(i)
                lang = getattr(editor.highlighter, "language", "Text") if hasattr(editor, "highlighter") else "Text"
                ext = "html"
                safe = "".join([c if c.isalnum() or c in ("_", "-") else "_" for c in title]) or f"tab_{i}"
                fname = f"{safe}_{i}.{ext}"
                fpath = os.path.join(self.session_dir, fname)
                html = editor.document().toHtml()
                html = self._rewrite_html_img_src(html)
                with open(fpath, "w", encoding="utf-8") as wf:
                    wf.write(html)
                tabs.append({"title": title, "lang": lang, "file": fname})
            meta_path = os.path.join(self.session_dir, "session.json")
            with open(meta_path, "w", encoding="utf-8") as mf:
                json.dump({"tabs": tabs}, mf, indent=2)
        except Exception:
            pass

    def _rewrite_html_img_src(self, html):
        try:
            os.makedirs(self.images_dir, exist_ok=True)
            import re
            from urllib.parse import urlparse, unquote
            def repl(m):
                src = m.group(1)
                if src.startswith('data:'):
                    return m.group(0)
                p = src
                try:
                    u = urlparse(src)
                    if u.scheme in ('file', ''):
                        p = unquote(u.path) if u.scheme == 'file' else src
                except Exception:
                    p = src
                if not os.path.isabs(p):
                    p = src
                if not os.path.exists(p):
                    return m.group(0)
                ext = os.path.splitext(p)[1] or '.png'
                dest = os.path.join(self.images_dir, f"img_{uuid.uuid4().hex}{ext}")
                try:
                    shutil.copy2(p, dest)
                    return f'src="{dest}"'
                except Exception:
                    return m.group(0)
            return re.sub(r'src="([^"]+)"', repl, html)
        except Exception:
            return html
        
    def create_menus(self):
        menubar = self.menuBar()
        
        # File Menu
        file_menu = menubar.addMenu("&File")
        
        new_tab_action = QAction("New &Tab", self)
        new_tab_action.setShortcut("Ctrl+T")
        new_tab_action.triggered.connect(self.new_tab)
        file_menu.addAction(new_tab_action)
        
        new_action = QAction("&New", self)
        new_action.setShortcut(QKeySequence.StandardKey.New)
        new_action.triggered.connect(self.new_file)
        file_menu.addAction(new_action)
        
        open_action = QAction("&Open", self)
        open_action.setShortcut(QKeySequence.StandardKey.Open)
        open_action.triggered.connect(self.open_file)
        file_menu.addAction(open_action)
        
        save_action = QAction("&Save", self)
        save_action.setShortcut(QKeySequence.StandardKey.Save)
        save_action.triggered.connect(self.save_file)
        file_menu.addAction(save_action)
        
        save_as_action = QAction("Save &As", self)
        save_as_action.setShortcut(QKeySequence.StandardKey.SaveAs)
        save_as_action.triggered.connect(self.save_file_as)
        file_menu.addAction(save_as_action)
        
        file_menu.addSeparator()
        
        print_action = QAction("&Print", self)
        print_action.setShortcut(QKeySequence.StandardKey.Print)
        print_action.triggered.connect(self.print_document)
        file_menu.addAction(print_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("E&xit", self)
        exit_action.setShortcut(QKeySequence.StandardKey.Quit)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Edit Menu
        edit_menu = menubar.addMenu("&Edit")
        
        undo_action = QAction("&Undo", self)
        undo_action.setShortcut(QKeySequence.StandardKey.Undo)
        undo_action.triggered.connect(lambda: self.current_tab().undo())
        edit_menu.addAction(undo_action)
        
        redo_action = QAction("&Redo", self)
        redo_action.setShortcut(QKeySequence.StandardKey.Redo)
        redo_action.triggered.connect(lambda: self.current_tab().redo())
        edit_menu.addAction(redo_action)
        
        edit_menu.addSeparator()
        
        cut_action = QAction("Cu&t", self)
        cut_action.setShortcut(QKeySequence.StandardKey.Cut)
        cut_action.triggered.connect(lambda: self.current_tab().cut())
        edit_menu.addAction(cut_action)
        
        copy_action = QAction("&Copy", self)
        copy_action.setShortcut(QKeySequence.StandardKey.Copy)
        copy_action.triggered.connect(lambda: self.current_tab().copy())
        edit_menu.addAction(copy_action)
        
        paste_action = QAction("&Paste", self)
        paste_action.setShortcut(QKeySequence.StandardKey.Paste)
        paste_action.triggered.connect(lambda: self.current_tab().paste())
        edit_menu.addAction(paste_action)
        
        edit_menu.addSeparator()
        
        select_all_action = QAction("Select &All", self)
        select_all_action.setShortcut(QKeySequence.StandardKey.SelectAll)
        select_all_action.triggered.connect(lambda: self.current_tab().selectAll())
        edit_menu.addAction(select_all_action)
        
        edit_menu.addSeparator()
        
        find_replace_action = QAction("&Find and Replace", self)
        find_replace_action.setShortcut("Ctrl+H")
        find_replace_action.triggered.connect(self.open_find_replace)
        edit_menu.addAction(find_replace_action)
        
        edit_menu.addSeparator()
        
        timestamp_action = QAction("Insert &Timestamp", self)
        timestamp_action.setShortcut("F5")
        timestamp_action.triggered.connect(self.insert_timestamp)
        edit_menu.addAction(timestamp_action)
        
        # Format Menu
        format_menu = menubar.addMenu("F&ormat")
        
        font_action = QAction("&Font", self)
        font_action.triggered.connect(self.change_font)
        format_menu.addAction(font_action)
        
        color_action = QAction("Text &Color", self)
        color_action.triggered.connect(self.change_color)
        format_menu.addAction(color_action)
        
        bg_color_action = QAction("&Background Color", self)
        bg_color_action.triggered.connect(self.change_bg_color)
        format_menu.addAction(bg_color_action)
        
        highlight_action = QAction("&Highlight Text", self)
        highlight_action.triggered.connect(self.highlight_text)
        format_menu.addAction(highlight_action)
        
        format_menu.addSeparator()
        
        bold_action = QAction("&Bold", self)
        bold_action.setShortcut(QKeySequence.StandardKey.Bold)
        bold_action.setCheckable(True)
        bold_action.triggered.connect(self.toggle_bold)
        format_menu.addAction(bold_action)
        
        italic_action = QAction("&Italic", self)
        italic_action.setShortcut(QKeySequence.StandardKey.Italic)
        italic_action.setCheckable(True)
        italic_action.triggered.connect(self.toggle_italic)
        format_menu.addAction(italic_action)
        
        underline_action = QAction("&Underline", self)
        underline_action.setShortcut(QKeySequence.StandardKey.Underline)
        underline_action.setCheckable(True)
        underline_action.triggered.connect(self.toggle_underline)
        format_menu.addAction(underline_action)
        
        strikethrough_action = QAction("&Strikethrough", self)
        strikethrough_action.triggered.connect(self.toggle_strikethrough)
        format_menu.addAction(strikethrough_action)
        
        format_menu.addSeparator()
        
        align_left_action = QAction("Align &Left", self)
        align_left_action.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignLeft))
        format_menu.addAction(align_left_action)
        
        align_center_action = QAction("Align &Center", self)
        align_center_action.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignCenter))
        format_menu.addAction(align_center_action)
        
        align_right_action = QAction("Align &Right", self)
        align_right_action.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignRight))
        format_menu.addAction(align_right_action)
        
        align_justify_action = QAction("&Justify", self)
        align_justify_action.triggered.connect(lambda: self.set_alignment(Qt.AlignmentFlag.AlignJustify))
        format_menu.addAction(align_justify_action)
        
        format_menu.addSeparator()
        
        increase_indent_action = QAction("Increase Indent", self)
        increase_indent_action.setShortcut("Ctrl+]")
        increase_indent_action.triggered.connect(self.increase_indent)
        format_menu.addAction(increase_indent_action)
        
        decrease_indent_action = QAction("Decrease Indent", self)
        decrease_indent_action.setShortcut("Ctrl+[")
        decrease_indent_action.triggered.connect(self.decrease_indent)
        format_menu.addAction(decrease_indent_action)
        
        # Insert Menu
        insert_menu = menubar.addMenu("&Insert")
        
        insert_image_action = QAction("Insert &Image", self)
        insert_image_action.triggered.connect(self.insert_image)
        insert_menu.addAction(insert_image_action)
        
        insert_table_action = QAction("Insert &Table", self)
        insert_table_action.triggered.connect(self.insert_table)
        insert_menu.addAction(insert_table_action)
        insert_ascii_table_action = QAction("Insert Table with |", self)
        insert_ascii_table_action.triggered.connect(self.insert_ascii_table)
        insert_menu.addAction(insert_ascii_table_action)
        
        insert_checkbox_action = QAction("Insert &Checkbox", self)
        insert_checkbox_action.triggered.connect(self.insert_checkbox)
        insert_menu.addAction(insert_checkbox_action)
        
        insert_menu.addSeparator()
        
        bullet_list_action = QAction("&Bullet List", self)
        bullet_list_action.triggered.connect(self.insert_bullet_list)
        insert_menu.addAction(bullet_list_action)
        
        numbered_list_action = QAction("&Numbered List", self)
        numbered_list_action.triggered.connect(self.insert_numbered_list)
        insert_menu.addAction(numbered_list_action)
        
        insert_menu.addSeparator()
        
        horizontal_line_action = QAction("Horizontal &Line", self)
        horizontal_line_action.triggered.connect(self.insert_horizontal_line)
        insert_menu.addAction(horizontal_line_action)
        
        insert_link_action = QAction("Insert Lin&k", self)
        insert_link_action.setShortcut("Ctrl+K")
        insert_link_action.triggered.connect(self.insert_link)
        insert_menu.addAction(insert_link_action)
        
        table_menu = menubar.addMenu("&Table")
        row_above_action = QAction("Insert Row Above", self)
        row_above_action.triggered.connect(self.table_insert_row_above)
        table_menu.addAction(row_above_action)
        row_below_action = QAction("Insert Row Below", self)
        row_below_action.triggered.connect(self.table_insert_row_below)
        table_menu.addAction(row_below_action)
        del_row_action = QAction("Delete Row", self)
        del_row_action.triggered.connect(self.table_delete_row)
        table_menu.addAction(del_row_action)
        table_menu.addSeparator()
        col_left_action = QAction("Insert Column Left", self)
        col_left_action.triggered.connect(self.table_insert_col_left)
        table_menu.addAction(col_left_action)
        col_right_action = QAction("Insert Column Right", self)
        col_right_action.triggered.connect(self.table_insert_col_right)
        table_menu.addAction(col_right_action)
        del_col_action = QAction("Delete Column", self)
        del_col_action.triggered.connect(self.table_delete_col)
        table_menu.addAction(del_col_action)
        table_menu.addSeparator()
        autofit_action = QAction("Auto-fit Columns", self)
        autofit_action.triggered.connect(self.table_autofit_columns)
        table_menu.addAction(autofit_action)
        
        # View Menu
        view_menu = menubar.addMenu("&View")
        
        self.stay_on_top_action = QAction("Always on &Top", self)
        self.stay_on_top_action.setCheckable(True)
        self.stay_on_top_action.triggered.connect(self.toggle_always_on_top)
        view_menu.addAction(self.stay_on_top_action)
        
        fullscreen_action = QAction("&Fullscreen", self)
        fullscreen_action.setShortcut("F11")
        fullscreen_action.triggered.connect(self.toggle_fullscreen)
        view_menu.addAction(fullscreen_action)
        
        view_menu.addSeparator()
        
        word_wrap_action = QAction("&Word Wrap", self)
        word_wrap_action.setCheckable(True)
        word_wrap_action.setChecked(True)
        word_wrap_action.triggered.connect(self.toggle_word_wrap)
        view_menu.addAction(word_wrap_action)
        
        view_menu.addSeparator()
        
        zoom_in_action = QAction("Zoom &In", self)
        zoom_in_action.setShortcut("Ctrl++")
        zoom_in_action.triggered.connect(self.zoom_in)
        view_menu.addAction(zoom_in_action)
        
        zoom_out_action = QAction("Zoom &Out", self)
        zoom_out_action.setShortcut("Ctrl+-")
        zoom_out_action.triggered.connect(self.zoom_out)
        view_menu.addAction(zoom_out_action)
        
        reset_zoom_action = QAction("&Reset Zoom", self)
        reset_zoom_action.setShortcut("Ctrl+0")
        reset_zoom_action.triggered.connect(self.reset_zoom)
        view_menu.addAction(reset_zoom_action)
        
        view_menu.addSeparator()
        
        # Theme submenu
        theme_menu = view_menu.addMenu("&Theme")
        for theme_name in self.themes.keys():
            theme_action = QAction(theme_name, self)
            theme_action.triggered.connect(lambda checked, t=theme_name: self.apply_theme(t))
            theme_menu.addAction(theme_action)
        
        view_menu.addSeparator()
        line_numbers_action = QAction("Show Line Numbers", self)
        line_numbers_action.setCheckable(True)
        line_numbers_action.setChecked(True)
        line_numbers_action.triggered.connect(self.toggle_line_numbers)
        view_menu.addAction(line_numbers_action)
        language_menu = view_menu.addMenu("&Language")
        for lang in ["Text","Python","R"]:
            act = QAction(lang, self)
            act.triggered.connect(lambda checked, l=lang: self.set_language(l))
            language_menu.addAction(act)
        
        # Tools Menu
        tools_menu = menubar.addMenu("&Tools")
        
        word_count_action = QAction("&Word Count", self)
        word_count_action.triggered.connect(self.show_word_count)
        tools_menu.addAction(word_count_action)
        
        char_count_action = QAction("&Character Count", self)
        char_count_action.triggered.connect(self.show_char_count)
        tools_menu.addAction(char_count_action)
        
        tools_menu.addSeparator()
        
        upper_case_action = QAction("Convert to &UPPERCASE", self)
        upper_case_action.triggered.connect(self.convert_to_upper)
        tools_menu.addAction(upper_case_action)
        
        lower_case_action = QAction("Convert to &lowercase", self)
        lower_case_action.triggered.connect(self.convert_to_lower)
        tools_menu.addAction(lower_case_action)
        
        title_case_action = QAction("Convert to &Title Case", self)
        title_case_action.triggered.connect(self.convert_to_title)
        tools_menu.addAction(title_case_action)
        
        tools_menu.addSeparator()
        
        remove_duplicates_action = QAction("Remove Duplicate &Lines", self)
        remove_duplicates_action.triggered.connect(self.remove_duplicate_lines)
        tools_menu.addAction(remove_duplicates_action)
        
        sort_lines_action = QAction("&Sort Lines A-Z", self)
        sort_lines_action.triggered.connect(self.sort_lines)
        tools_menu.addAction(sort_lines_action)
        
        tools_menu.addSeparator()
        correct_action = QAction("&Correct Selection", self)
        correct_action.triggered.connect(self.correct_selection)
        tools_menu.addAction(correct_action)
        tools_menu.addSeparator()
        run_menu_action = QAction("▶ &Run", self)
        run_menu_action.triggered.connect(self.run_current)
        tools_menu.addAction(run_menu_action)
        run_sel_menu_action = QAction("▶ Run &Selection", self)
        run_sel_menu_action.triggered.connect(self.run_selection)
        tools_menu.addAction(run_sel_menu_action)
        samples_action = QAction("Create &Sample Files", self)
        samples_action.triggered.connect(self.create_sample_files)
        tools_menu.addAction(samples_action)
        
    def create_toolbar(self):
        toolbar = QToolBar()
        toolbar.setMovable(False)
        self.addToolBar(toolbar)
        
        # File operations
        new_btn = QAction("📄 New", self)
        new_btn.triggered.connect(self.new_file)
        toolbar.addAction(new_btn)
        
        open_btn = QAction("📂 Open", self)
        open_btn.triggered.connect(self.open_file)
        toolbar.addAction(open_btn)
        
        save_btn = QAction("💾 Save", self)
        save_btn.triggered.connect(self.save_file)
        toolbar.addAction(save_btn)
        
        toolbar.addSeparator()
        
        # Format operations
        bold_btn = QAction("B", self)
        bold_btn.triggered.connect(self.toggle_bold)
        toolbar.addAction(bold_btn)
        
        italic_btn = QAction("I", self)
        italic_btn.triggered.connect(self.toggle_italic)
        toolbar.addAction(italic_btn)
        
        underline_btn = QAction("U", self)
        underline_btn.triggered.connect(self.toggle_underline)
        toolbar.addAction(underline_btn)
        
        toolbar.addSeparator()
        
        # Insert operations
        image_btn = QAction("🖼️ Image", self)
        image_btn.triggered.connect(self.insert_image)
        toolbar.addAction(image_btn)
        
        table_btn = QAction("📊 Table", self)
        table_btn.triggered.connect(self.insert_table)
        toolbar.addAction(table_btn)
        
        checkbox_btn = QAction("☐ Check", self)
        checkbox_btn.triggered.connect(self.insert_checkbox)
        toolbar.addAction(checkbox_btn)
        
        toolbar.addSeparator()
        
        # Zoom controls
        zoom_out_btn = QAction("🔍-", self)
        zoom_out_btn.triggered.connect(self.zoom_out)
        toolbar.addAction(zoom_out_btn)
        
        zoom_in_btn = QAction("🔍+", self)
        zoom_in_btn.triggered.connect(self.zoom_in)
        toolbar.addAction(zoom_in_btn)
        toolbar.addSeparator()
        run_btn = QAction("▶ Run", self)
        run_btn.triggered.connect(self.run_current)
        toolbar.addAction(run_btn)
        run_sel_btn = QAction("▶ Run Selection", self)
        run_sel_btn.triggered.connect(self.run_selection)
        toolbar.addAction(run_sel_btn)
        self.screen_combo = QComboBox()
        self.screen_combo.addItem("All Screens")
        for s in QApplication.screens():
            self.screen_combo.addItem(s.name())
        toolbar.addWidget(self.screen_combo)
        start_rec_btn = QAction("⏺ Start", self)
        start_rec_btn.triggered.connect(self.start_screen_recording)
        toolbar.addAction(start_rec_btn)
        stop_rec_btn = QAction("⏹ Stop", self)
        stop_rec_btn.triggered.connect(self.stop_screen_recording)
        toolbar.addAction(stop_rec_btn)
        capture_now_btn = QAction("📷 Capture Now", self)
        capture_now_btn.triggered.connect(self.capture_now)
        toolbar.addAction(capture_now_btn)
        stay_on_top_btn = QAction("📌 On Top", self)
        stay_on_top_btn.setCheckable(True)
        def _top_clicked(checked):
            try:
                self.toggle_always_on_top(checked)
                if hasattr(self, 'stay_on_top_action'):
                    self.stay_on_top_action.setChecked(checked)
            except Exception:
                pass
        stay_on_top_btn.triggered.connect(_top_clicked)
        if hasattr(self, 'stay_on_top_action'):
            self.stay_on_top_action.toggled.connect(stay_on_top_btn.setChecked)
        toolbar.addAction(stay_on_top_btn)
        
    def text_changed(self):
        self.is_modified = True
        self.update_title()
        try:
            editor = self.current_tab()
            lang = getattr(editor.highlighter, 'language', 'Text') if hasattr(editor, 'highlighter') else 'Text'
            if lang == 'Python':
                with open(self.python_temp_path, 'w', encoding='utf-8') as f:
                    f.write(editor.toPlainText())
        except Exception:
            pass
        
    def update_title(self):
        title = "Ultra Advanced Notepad++"
        if self.current_file:
            title += f" - {self.current_file}"
        if self.is_modified:
            title += " *"
        self.setWindowTitle(title)
        
    def update_status(self):
        text_edit = self.current_tab()
        if text_edit:
            cursor = text_edit.textCursor()
            line = cursor.blockNumber() + 1
            col = cursor.columnNumber() + 1
            chars = len(text_edit.toPlainText())
            words = len(text_edit.toPlainText().split())
            self.status_bar.showMessage(
                f"Line: {line}, Col: {col} | Words: {words} | Characters: {chars} | Zoom: {self.zoom_level}%"
            )
        
    def new_file(self):
        self.new_tab()
            
    def open_file(self):
        filename, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", 
            "All Files (*.*)"
        )
        if not filename:
            return False
        try:
            editor = self.new_tab()
            lower = filename.lower()
            if lower.endswith(('.csv', '.tsv', '.txt')):
                tmp = os.path.join(self.session_dir, f"duckdb_{os.path.basename(filename)}.pipe")
                os.makedirs(self.session_dir, exist_ok=True)
                try:
                    if duckdb:
                        con = duckdb.connect()
                        con.execute(f"COPY (SELECT * FROM read_csv_auto('{filename}')) TO '{tmp}' (DELIMITER '|', HEADER TRUE)")
                    else:
                        raise Exception('duckdb not available')
                except Exception:
                    import csv
                    with open(filename, 'r', encoding='utf-8', errors='replace', newline='') as fin, open(tmp, 'w', encoding='utf-8', newline='') as fout:
                        sample = fin.read(4096)
                        fin.seek(0)
                        try:
                            dialect = csv.Sniffer().sniff(sample)
                        except Exception:
                            dialect = csv.excel
                        reader = csv.reader(fin, dialect)
                        w = csv.writer(fout, delimiter='|')
                        for row in reader:
                            w.writerow(row)
                aligned = self._align_pipe_file(tmp)
                with open(aligned, 'r', encoding='utf-8', errors='replace') as f:
                    editor.setPlainText(f.read())
                try:
                    os.remove(tmp)
                    os.remove(aligned)
                except Exception:
                    pass
            elif lower.endswith(('.parquet', '.parq')):
                tmp = os.path.join(self.session_dir, f"duckdb_{os.path.basename(filename)}.pipe")
                os.makedirs(self.session_dir, exist_ok=True)
                try:
                    if duckdb:
                        con = duckdb.connect()
                        con.execute(f"COPY (SELECT * FROM read_parquet('{filename}')) TO '{tmp}' (DELIMITER '|', HEADER TRUE)")
                    elif pl:
                        df = pl.read_parquet(filename)
                        df.write_csv(tmp, separator='|')
                    else:
                        raise Exception('No parquet loader available')
                except Exception as e:
                    QMessageBox.critical(self, 'Open', f'Failed to load parquet:\n{e}')
                    self.tab_widget.removeTab(self.tab_widget.currentIndex())
                    return False
                aligned = self._align_pipe_file(tmp)
                with open(aligned, 'r', encoding='utf-8', errors='replace') as f:
                    editor.setPlainText(f.read())
                try:
                    os.remove(tmp)
                    os.remove(aligned)
                except Exception:
                    pass
            elif lower.endswith(('.xlsx', '.xls')) and pl:
                df = pl.read_excel(filename)
                tmp = os.path.join(self.session_dir, f"polars_{os.path.basename(filename)}.pipe")
                os.makedirs(self.session_dir, exist_ok=True)
                df.write_csv(tmp, separator='|')
                aligned = self._align_pipe_file(tmp)
                with open(aligned, 'r', encoding='utf-8', errors='replace') as f:
                    editor.setPlainText(f.read())
                try:
                    os.remove(tmp)
                    os.remove(aligned)
                except Exception:
                    pass
            elif lower.endswith(('.json', '.ndjson', '.jsonl')):
                with open(filename, 'r', encoding='utf-8', errors='replace') as f:
                    raw = f.read()
                try:
                    obj = json.loads(raw)
                    editor.setPlainText(json.dumps(obj, indent=2, ensure_ascii=False))
                except Exception:
                    lines = []
                    for line in raw.splitlines():
                        try:
                            lines.append(json.dumps(json.loads(line), indent=2, ensure_ascii=False))
                        except Exception:
                            lines.append(line)
                    editor.setPlainText("\n".join(lines))
            elif lower.endswith(('.pdf')) and PdfReader:
                reader = PdfReader(filename)
                pages = []
                for p in reader.pages:
                    pages.append(p.extract_text() or '')
                editor.setPlainText("\n\n".join(pages))
            else:
                with open(filename, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                if lower.endswith(('.html', '.htm')):
                    editor.setHtml(content)
                else:
                    editor.setPlainText(content)
            if lower.endswith('.py') and hasattr(editor, 'setLanguage'):
                editor.setLanguage('Python')
            elif lower.endswith('.r') and hasattr(editor, 'setLanguage'):
                editor.setLanguage('R')
            elif lower.endswith('.md') and hasattr(editor, 'setLanguage'):
                editor.setLanguage('Text')
            self.current_file = filename
            self.is_modified = False
            self.tab_widget.setTabText(self.tab_widget.currentIndex(), os.path.basename(filename))
            self.update_title()
            self.status_bar.showMessage("File opened successfully", 3000)
            return True
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not open file:\n{e}")
            return False

    def _align_pipe_file(self, src_path):
        try:
            dst = src_path + ".mdtbl"
            with open(src_path, 'r', encoding='utf-8', errors='replace') as f:
                lines = f.read().splitlines()
            if not lines:
                with open(dst, 'w', encoding='utf-8') as out:
                    out.write('')
                return dst
            header = lines[0].split('|')
            widths = [len(h) for h in header]
            sample_count = min(len(lines), 1000)
            for i in range(1, sample_count):
                cells = lines[i].split('|')
                for j in range(min(len(widths), len(cells))):
                    w = len(cells[j])
                    if w > widths[j]:
                        widths[j] = w
            def fmt_row(cells):
                padded = []
                for j in range(len(widths)):
                    val = cells[j] if j < len(cells) else ''
                    padded.append(val.ljust(widths[j]))
                return '| ' + ' | '.join(padded) + ' |\n'
            with open(dst, 'w', encoding='utf-8') as out:
                out.write(fmt_row(header))
                out.write('| ' + ' | '.join(['-' * w for w in widths]) + ' |\n')
                for i in range(1, len(lines)):
                    out.write(fmt_row(lines[i].split('|')))
            return dst
        except Exception:
            return src_path
    
    def save_file(self):
        if self.current_file:
            return self.save_to_file(self.current_file)
        else:
            return self.save_file_as()
    
    def save_file_as(self):
        filename, selected_filter = QFileDialog.getSaveFileName(
            self, "Save File As", "", 
            "Text Files (*.txt);;Markdown Files (*.md);;HTML Files (*.html);;PDF Files (*.pdf);;Word Files (*.docx);;All Files (*.*)"
        )
        if filename:
            return self.save_to_file(filename)
        return False
    
    def save_to_file(self, filename):
        try:
            text_edit = self.current_tab()
            lower = filename.lower()
            if lower.endswith('.pdf'):
                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(filename)
                text_edit.document().print(printer)
            elif lower.endswith('.html'):
                html = text_edit.document().toHtml()
                base = os.path.splitext(os.path.basename(filename))[0]
                assets = os.path.join(os.path.dirname(filename), base + '_assets')
                os.makedirs(assets, exist_ok=True)
                html2 = self._rewrite_html_img_src_for_export(html, assets)
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(html2)
            elif lower.endswith('.docx'):
                try:
                    import docx
                    from docx.shared import Inches
                except Exception:
                    raise Exception('Microsoft Word export requires python-docx to be installed')
                document = docx.Document()
                html = text_edit.document().toHtml()
                self._export_docx_from_html(document, html)
                document.save(filename)
            else:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(text_edit.toPlainText())
            self.current_file = filename
            self.is_modified = False
            self.tab_widget.setTabText(self.tab_widget.currentIndex(), os.path.basename(filename))
            self.update_title()
            self.status_bar.showMessage("Saved", 3000)
            return True
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not save file:\n{e}")
            return False

    def _export_docx_from_html(self, document, html):
        try:
            import re
            from urllib.parse import urlparse, unquote
            import io
            import base64
            def _resolve_img_src(src):
                if src.startswith('data:'):
                    return ('data', src)
                try:
                    u = urlparse(src)
                    if u.scheme == 'file':
                        p = unquote(u.path)
                        if os.name == 'nt' and p.startswith('/') and len(p) > 3 and p[2] == ':':
                            p = p.lstrip('/')
                        return ('path', p)
                    elif u.scheme in ('', None):
                        return ('path', src)
                except Exception:
                    pass
                return ('path', src)
            def _add_picture(src):
                kind, val = _resolve_img_src(src)
                if kind == 'data':
                    try:
                        m = re.match(r'data:([^;]+);base64,(.+)', val)
                        if m:
                            data_b64 = m.group(2)
                            buf = io.BytesIO()
                            buf.write(base64.b64decode(data_b64))
                            buf.seek(0)
                            document.add_picture(buf)
                            return True
                    except Exception:
                        return False
                    return False
                p = val
                if not os.path.isabs(p):
                    p = val
                if os.path.exists(p):
                    try:
                        from docx.shared import Inches
                        document.add_picture(p, width=Inches(6))
                        return True
                    except Exception:
                        document.add_paragraph(f"[Image: {p}]")
                        return True
                return False
            # Extract screenshot-like tables and images
            tables = re.findall(r'<table[\s\S]*?</table>', html)
            for t in tables:
                imgs = re.findall(r'<img[^>]*src="([^"]+)"', t)
                cells = re.findall(r'<td[\s\S]*?>([\s\S]*?)</td>', t)
                notes_text = ''
                if len(cells) >= 2:
                    right = re.sub(r'<[^>]+>', '', cells[1])
                    notes_text = re.sub(r'\s+', ' ', right).strip()
                for src in imgs:
                    ok = _add_picture(src)
                    if ok:
                        if notes_text:
                            document.add_paragraph(notes_text)
                        document.add_paragraph("")
            # Remaining text content
            html_no_tables = re.sub(r'<table[\s\S]*?</table>', '', html)
            other_imgs = re.findall(r'<img[^>]*src="([^"]+)"', html_no_tables)
            for src in other_imgs:
                ok = _add_picture(src)
                if ok:
                    document.add_paragraph("")
            text = re.sub(r'<br\s*/?>', '\n', html_no_tables)
            text = re.sub(r'<[^>]+>', '', text)
            lines = [l.strip() for l in text.splitlines()]
            for l in lines:
                if l:
                    document.add_paragraph(l)
        except Exception:
            # Fallback: plain text
            document.add_paragraph(re.sub(r'<[^>]+>', '', html))
    
    def auto_save(self):
        if self.is_modified and self.current_file:
            self.save_to_file(self.current_file)
            
    def maybe_save(self):
        if self.is_modified:
            reply = QMessageBox.question(
                self, "Save Changes?",
                "Do you want to save changes to the document?",
                QMessageBox.StandardButton.Save | 
                QMessageBox.StandardButton.Discard | 
                QMessageBox.StandardButton.Cancel
            )
            if reply == QMessageBox.StandardButton.Save:
                return self.save_file()
            elif reply == QMessageBox.StandardButton.Cancel:
                return False
        return True
        
    def closeEvent(self, event):
        self.save_session()
        event.accept()
            
    def print_document(self):
        printer = QPrinter()
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.current_tab().print(printer)
            
    def change_font(self):
        font, ok = QFontDialog.getFont(self.current_tab().currentFont(), self)
        if ok:
            self.current_tab().setCurrentFont(font)
            
    def change_color(self):
        color = QColorDialog.getColor(self.current_tab().textColor(), self)
        if color.isValid():
            self.current_tab().setTextColor(color)
            
    def change_bg_color(self):
        color = QColorDialog.getColor(self.current_tab().textBackgroundColor(), self)
        if color.isValid():
            self.current_tab().setTextBackgroundColor(color)
    
    def highlight_text(self):
        color = QColorDialog.getColor(Qt.GlobalColor.yellow, self)
        if color.isValid():
            cursor = self.current_tab().textCursor()
            if cursor.hasSelection():
                fmt = QTextCharFormat()
                fmt.setBackground(color)
                cursor.mergeCharFormat(fmt)
            
    def toggle_bold(self):
        fmt = self.current_tab().currentCharFormat()
        weight = QFont.Weight.Bold if fmt.fontWeight() != QFont.Weight.Bold else QFont.Weight.Normal
        fmt.setFontWeight(weight)
        self.current_tab().setCurrentCharFormat(fmt)
        
    def toggle_italic(self):
        fmt = self.current_tab().currentCharFormat()
        fmt.setFontItalic(not fmt.fontItalic())
        self.current_tab().setCurrentCharFormat(fmt)
        
    def toggle_underline(self):
        fmt = self.current_tab().currentCharFormat()
        fmt.setFontUnderline(not fmt.fontUnderline())
        self.current_tab().setCurrentCharFormat(fmt)
    
    def toggle_strikethrough(self):
        fmt = self.current_tab().currentCharFormat()
        fmt.setFontStrikeOut(not fmt.fontStrikeOut())
        self.current_tab().setCurrentCharFormat(fmt)
    
    def set_alignment(self, alignment):
        self.current_tab().setAlignment(alignment)
    
    def increase_indent(self):
        cursor = self.current_tab().textCursor()
        block_format = cursor.blockFormat()
        block_format.setIndent(block_format.indent() + 1)
        cursor.setBlockFormat(block_format)
    
    def decrease_indent(self):
        cursor = self.current_tab().textCursor()
        block_format = cursor.blockFormat()
        indent = max(0, block_format.indent() - 1)
        block_format.setIndent(indent)
        cursor.setBlockFormat(block_format)
        
    def open_find_replace(self):
        dialog = FindReplaceDialog(self)
        dialog.show()
    
    def insert_timestamp(self):
        cursor = self.current_tab().textCursor()
        timestamp = QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss")
        cursor.insertText(timestamp)
        
    def insert_image(self):
        filename, _ = QFileDialog.getOpenFileName(
            self, "Insert Image", "", 
            "Images (*.png *.jpg *.jpeg *.bmp *.gif);;All Files (*.*)"
        )
        if filename:
            cursor = self.current_tab().textCursor()
            image = QImage(filename)
            if not image.isNull():
                max_width = 600
                if image.width() > max_width:
                    image = image.scaledToWidth(max_width, Qt.TransformationMode.SmoothTransformation)
                
                image_format = QTextImageFormat()
                image_format.setName(filename)
                image_format.setWidth(image.width())
                image_format.setHeight(image.height())
                cursor.insertImage(image_format)
            else:
                QMessageBox.warning(self, "Error", "Could not load image")
                
    def insert_table(self):
        dialog = TableDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            rows = dialog.rows_spin.value()
            cols = dialog.cols_spin.value()
            
            cursor = self.current_tab().textCursor()
            table_format = QTextTableFormat()
            table_format.setBorder(2)
            table_format.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Solid)
            table_format.setCellPadding(4)
            table_format.setCellSpacing(0)
            table_format.setTopMargin(4)
            table_format.setBottomMargin(12)
            text_color = self.current_tab().palette().color(QPalette.ColorRole.Text)
            table_format.setBorderBrush(text_color)
            cursor.insertTable(rows, cols, table_format)
            table = cursor.currentTable()
            if table:
                self._format_table_grid(table)

    def insert_ascii_table(self):
        dialog = TableDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            rows = dialog.rows_spin.value()
            cols = dialog.cols_spin.value()
            header = "|" + "|".join(["   " for _ in range(cols)]) + "|\n"
            divider = "|" + "|".join(["---" for _ in range(cols)]) + "|\n"
            body_lines = []
            for _ in range(max(1, rows)):
                body_lines.append("|" + "|".join(["   " for _ in range(cols)]) + "|")
            text = header + divider + "\n".join(body_lines) + "\n"
            self.current_tab().textCursor().insertText(text)
            
    def table_insert_row_above(self):
        cursor = self.current_tab().textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            if cell.isValid():
                table.insertRows(cell.row(), 1)
                self._format_table_grid(table)
            else:
                QMessageBox.information(self, "Table", "Place the cursor inside a table cell")
        else:
            QMessageBox.information(self, "Table", "Place the cursor inside a table")
            
    def table_insert_row_below(self):
        cursor = self.current_tab().textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            if cell.isValid():
                table.insertRows(cell.row() + 1, 1)
                self._format_table_grid(table)
            else:
                QMessageBox.information(self, "Table", "Place the cursor inside a table cell")
        else:
            QMessageBox.information(self, "Table", "Place the cursor inside a table")
            
    def table_delete_row(self):
        cursor = self.current_tab().textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            if cell.isValid():
                table.removeRows(cell.row(), 1)
                self._format_table_grid(table)
            else:
                QMessageBox.information(self, "Table", "Place the cursor inside a table cell")
        else:
            QMessageBox.information(self, "Table", "Place the cursor inside a table")
            
    def table_insert_col_left(self):
        cursor = self.current_tab().textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            if cell.isValid():
                table.insertColumns(cell.column(), 1)
                self._format_table_grid(table)
            else:
                QMessageBox.information(self, "Table", "Place the cursor inside a table cell")
        else:
            QMessageBox.information(self, "Table", "Place the cursor inside a table")
            
    def table_insert_col_right(self):
        cursor = self.current_tab().textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            if cell.isValid():
                table.insertColumns(cell.column() + 1, 1)
                self._format_table_grid(table)
            else:
                QMessageBox.information(self, "Table", "Place the cursor inside a table cell")
        else:
            QMessageBox.information(self, "Table", "Place the cursor inside a table")
            
    def table_delete_col(self):
        cursor = self.current_tab().textCursor()
        table = cursor.currentTable()
        if table:
            cell = table.cellAt(cursor)
            if cell.isValid():
                table.removeColumns(cell.column(), 1)
                self._format_table_grid(table)
            else:
                QMessageBox.information(self, "Table", "Place the cursor inside a table cell")
        else:
            QMessageBox.information(self, "Table", "Place the cursor inside a table")
            
    def table_autofit_columns(self):
        text_edit = self.current_tab()
        cursor = text_edit.textCursor()
        table = cursor.currentTable()
        if not table:
            QMessageBox.information(self, "Table", "Place the cursor inside a table")
            return
        rows = table.rows()
        cols = table.columns()
        fm = text_edit.fontMetrics()
        col_widths = [0] * cols
        for c in range(cols):
            maxw = 0
            for r in range(rows):
                cell = table.cellAt(r, c)
                c1 = cell.firstCursorPosition()
                c2 = cell.lastCursorPosition()
                tmp = c1
                tmp.setPosition(c2.position(), QTextCursor.MoveMode.KeepAnchor)
                s = tmp.selectedText()
                w = fm.horizontalAdvance(s) + 16
                if w > maxw:
                    maxw = w
            col_widths[c] = maxw
        total = float(sum(col_widths)) if sum(col_widths) > 0 else float(cols)
        constraints = []
        for w in col_widths:
            percent = max(5.0, (w / total) * 100.0)
            constraints.append(QTextLength(QTextLength.Type.PercentageLength, percent))
        fmt = table.format()
        fmt.setColumnWidthConstraints(constraints)
        table.setFormat(fmt)
        self._format_table_grid(table)

    def _format_table_grid(self, table):
        text_color = self.current_tab().palette().color(QPalette.ColorRole.Text)
        fmt = table.format()
        fmt.setBorder(2)
        fmt.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Solid)
        fmt.setCellPadding(4)
        fmt.setCellSpacing(0)
        fmt.setTopMargin(4)
        fmt.setBottomMargin(12)
        fmt.setBorderBrush(text_color)
        table.setFormat(fmt)
        cell_fmt = QTextTableCellFormat()
        cell_fmt.setBorder(1)
        cell_fmt.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Solid)
        cell_fmt.setBorderBrush(text_color)
        rows = table.rows()
        cols = table.columns()
        for r in range(rows):
            for c in range(cols):
                cell = table.cellAt(r, c)
                cell.setFormat(cell_fmt)
            
    def insert_checkbox(self):
        cursor = self.current_tab().textCursor()
        cursor.insertText("☐ ")
            
    def insert_bullet_list(self):
        cursor = self.current_tab().textCursor()
        list_format = QTextListFormat()
        list_format.setStyle(QTextListFormat.Style.ListDisc)
        cursor.createList(list_format)
        
    def insert_numbered_list(self):
        cursor = self.current_tab().textCursor()
        list_format = QTextListFormat()
        list_format.setStyle(QTextListFormat.Style.ListDecimal)
        cursor.createList(list_format)
    
    def insert_horizontal_line(self):
        cursor = self.current_tab().textCursor()
        cursor.insertHtml("<hr>")
    
    def insert_link(self):
        text, ok1 = QInputDialog.getText(self, "Insert Link", "Link text:")
        if ok1 and text:
            url, ok2 = QInputDialog.getText(self, "Insert Link", "URL:")
            if ok2 and url:
                cursor = self.current_tab().textCursor()
                cursor.insertHtml(f'<a href="{url}">{text}</a>')

    def open_screen_capture(self):
        try:
            dlg = QDialog(self)
            dlg.setWindowTitle("Capture Screens")
            layout = QVBoxLayout(dlg)
            screens = QApplication.screens()
            for idx, s in enumerate(screens):
                geo = s.geometry()
                btn = QPushButton(f"Capture Screen {idx+1}: {s.name()} ({geo.width()}x{geo.height()})")
                def make_cb(i):
                    return lambda: (self.capture_screen_by_index(i), dlg.accept())
                btn.clicked.connect(make_cb(idx))
                layout.addWidget(btn)
            all_btn = QPushButton("Capture All Screens")
            all_btn.clicked.connect(lambda: (self.capture_all_screens(), dlg.accept()))
            layout.addWidget(all_btn)
            cancel_btn = QPushButton("Cancel")
            cancel_btn.clicked.connect(dlg.reject)
            layout.addWidget(cancel_btn)
            dlg.exec()
        except Exception:
            pass

    def capture_screen_by_index(self, idx):
        try:
            screens = QApplication.screens()
            if 0 <= idx < len(screens):
                s = screens[idx]
                pix = s.grabWindow(0)
                self._insert_screenshot_pixmap(pix)
                self.status_bar.showMessage("Captured screen", 3000)
        except Exception:
            pass

    def capture_all_screens(self):
        try:
            screens = QApplication.screens()
            for s in screens:
                pix = s.grabWindow(0)
                self._insert_screenshot_pixmap(pix)
            self.status_bar.showMessage("Captured all screens", 3000)
        except Exception:
            pass

    def _insert_screenshot_pixmap(self, pixmap):
        try:
            os.makedirs(self.images_dir, exist_ok=True)
            path = os.path.join(self.images_dir, f"screenshot_{uuid.uuid4().hex}.png")
            pixmap.save(path, 'PNG')
            editor = self.current_tab()
            img = QImage(path)
            view_w = editor.viewport().width()
            max_img_w = int(max(300, view_w * 0.6))
            disp_w = img.width()
            if disp_w > max_img_w:
                disp_w = max_img_w
            disp_h = int(img.height() * (disp_w / img.width())) if img.width() else img.height()
            cursor = editor.textCursor()
            tf = QTextTableFormat()
            tf.setBorder(1)
            tf.setCellPadding(8)
            tf.setCellSpacing(8)
            tf.setColumnWidthConstraints([
                QTextLength(QTextLength.PercentageLength, 62),
                QTextLength(QTextLength.PercentageLength, 38)
            ])
            table = cursor.insertTable(1, 2, tf)
            left = table.cellAt(0, 0).firstCursorPosition()
            imgfmt = QTextImageFormat()
            imgfmt.setName(path)
            imgfmt.setWidth(disp_w)
            imgfmt.setHeight(disp_h)
            left.insertImage(imgfmt)
            right = table.cellAt(0, 1).firstCursorPosition()
            bold = QTextCharFormat()
            bold.setFontWeight(QFont.Weight.Bold)
            right.mergeCharFormat(bold)
            right.insertText("Notes:\n")
            normal = QTextCharFormat()
            normal.setFontWeight(QFont.Weight.Normal)
            right.mergeCharFormat(normal)
            editor.ensureCursorVisible()
        except Exception:
            pass

    def _insert_image_with_note(self, path, note, screen_name=None, ts=None):
        try:
            editor = self.current_tab()
            img = QImage(path)
            view_w = editor.viewport().width()
            max_img_w = int(max(300, view_w * 0.6))
            disp_w = img.width()
            if disp_w > max_img_w:
                disp_w = max_img_w
            disp_h = int(img.height() * (disp_w / img.width())) if img.width() else img.height()
            cursor = editor.textCursor()
            try:
                cursor.movePosition(QTextCursor.MoveOperation.End)
                editor.setTextCursor(cursor)
            except Exception:
                pass
            from PyQt6.QtCore import QUrl
            url = QUrl.fromLocalFile(path).toString()
            info = (note or "").strip()
            if info:
                info_html = info.replace('\n', '<br/>')
                html = (
                    f"<table style='border:1px solid #ccc;border-collapse:collapse;margin:8px 0;'>"
                    f"<tr>"
                    f"<td style='padding:8px;vertical-align:top;'>"
                    f"<img src='{url}' width='{disp_w}' height='{disp_h}'/>"
                    f"</td>"
                    f"<td style='padding:8px;vertical-align:top;'>"
                    f"<b>Notes:</b><br/>{info_html}"
                    f"</td>"
                    f"</tr>"
                    f"</table><br/>"
                )
            else:
                html = (
                    f"<div style='margin:8px 0;'>"
                    f"<img src='{url}' width='{disp_w}' height='{disp_h}'/>"
                    f"</div><br/>"
                )
            cursor.insertHtml(html)
            editor.ensureCursorVisible()
            try:
                self.status_bar.showMessage("Inserted image", 2000)
            except Exception:
                pass
        except Exception:
            pass

    def _insert_header_note(self, note):
        try:
            editor = self.current_tab()
            cursor = editor.textCursor()
            try:
                cursor.movePosition(QTextCursor.MoveOperation.End)
                editor.setTextCursor(cursor)
            except Exception:
                pass
            info = (note or "").strip()
            if info:
                info_html = info.replace('\n', '<br/>')
                html = (
                    f"<div style='margin:8px 0;'>"
                    f"<b>Notes:</b><br/>{info_html}"
                    f"</div><br/>"
                )
                cursor.insertHtml(html)
                editor.ensureCursorVisible()
                try:
                    self.status_bar.showMessage("Inserted note", 2000)
                except Exception:
                    pass
        except Exception:
            pass

    def capture_now(self):
        try:
            choice = self.screen_combo.currentText() if hasattr(self, 'screen_combo') else "All Screens"
            now = QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss")
            if choice == "All Screens":
                for s in QApplication.screens():
                    pix = s.grabWindow(0)
                    os.makedirs(self.images_dir, exist_ok=True)
                    path = os.path.join(self.images_dir, f"screencap_{uuid.uuid4().hex}.png")
                    pix.save(path, 'PNG')
                    self.captured_images.append({"path": path, "screen": s.name(), "time": now})
                    self._insert_image_with_note(path, "", s.name(), now)
            else:
                target = None
                for s in QApplication.screens():
                    if s.name() == choice:
                        target = s
                        break
                if target:
                    pix = target.grabWindow(0)
                    os.makedirs(self.images_dir, exist_ok=True)
                    path = os.path.join(self.images_dir, f"screencap_{uuid.uuid4().hex}.png")
                    pix.save(path, 'PNG')
                    self.captured_images.append({"path": path, "screen": target.name(), "time": now})
                    self._insert_image_with_note(path, "", target.name(), now)
            self.status_bar.showMessage("Captured", 1500)
        except Exception:
            pass

    def start_screen_recording(self):
        try:
            self.captured_images = []
            self.capture_active = True
            self.capture_timer.start()
            if not self._capture_filter_installed:
                QApplication.instance().installEventFilter(self)
                self._capture_filter_installed = True
            self.status_bar.showMessage("Recording screens", 2000)
        except Exception:
            pass

    def stop_screen_recording(self):
        try:
            self.capture_timer.stop()
            self.capture_active = False
            if not self.captured_images:
                return
            if self._capture_filter_installed:
                try:
                    QApplication.instance().removeEventFilter(self)
                except Exception:
                    pass
                self._capture_filter_installed = False
            dlg = QDialog(self)
            dlg.setWindowTitle("Insert Captures")
            dlg.setWindowFlag(Qt.WindowType.WindowMaximizeButtonHint, True)
            dlg.setSizeGripEnabled(True)
            v = QVBoxLayout(dlg)
            select_all = QCheckBox("Select All")
            v.addWidget(select_all)
            scroll = QScrollArea(dlg)
            scroll.setWidgetResizable(True)
            container = QWidget()
            container_layout = QVBoxLayout(container)
            scroll.setWidget(container)
            v.addWidget(scroll)
            checks = []
            rows = []
            def open_preview(d, comm_widget):
                ndlg = QDialog(self)
                ndlg.setWindowTitle("Preview")
                ndlg.setWindowFlag(Qt.WindowType.WindowMaximizeButtonHint, True)
                ndlg.setSizeGripEnabled(True)
                lay = QVBoxLayout(ndlg)
                img_label = QLabel()
                pm2 = QPixmap(d.get('path'))
                if not pm2.isNull():
                    ww = max(500, int(self.width() * 0.7))
                    hh = int(pm2.height() * (ww / pm2.width())) if pm2.width() else pm2.height()
                    img_label.setPixmap(pm2.scaled(ww, hh, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
                lay.addWidget(img_label)
                comment_label = QLabel("Comment for this image:")
                lay.addWidget(comment_label)
                comment_edit = QPlainTextEdit()
                comment_edit.setPlainText(comm_widget.text())
                lay.addWidget(comment_edit)
                btns2 = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
                lay.addWidget(btns2)
                def on_ok2():
                    note = comment_edit.toPlainText().strip()
                    comm_widget.setText(note)
                    self._insert_image_with_note(d.get('path'), note, d.get('screen'), d.get('time'))
                    ndlg.accept()
                btns2.accepted.connect(on_ok2)
                btns2.rejected.connect(ndlg.reject)
                ndlg.exec()
            for item in list(self.captured_images):
                row_widget = QWidget()
                row = QHBoxLayout(row_widget)
                thumb = QLabel()
                pm = QPixmap(item.get('path'))
                if not pm.isNull():
                    w = 180
                    h = int(pm.height() * (w / pm.width())) if pm.width() else pm.height()
                    thumb.setPixmap(pm.scaled(w, h, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
                thumb.setCursor(Qt.CursorShape.PointingHandCursor)
                cb = QCheckBox(f"{item.get('screen')} @ {item.get('time')}")
                cb.setChecked(True)
                cb.item_data = item
                checks.append(cb)
                preview_btn = QPushButton("Preview")
                delete_btn = QPushButton("Delete")
                comm = QLineEdit()
                comm.setPlaceholderText("Comment")
                def on_preview():
                    open_preview(item, comm)
                preview_btn.clicked.connect(on_preview)
                def on_thumb_click(event):
                    open_preview(item, comm)
                thumb.mousePressEvent = on_thumb_click
                def on_delete():
                    try:
                        # remove from captured list
                        self.captured_images = [x for x in self.captured_images if x.get('path') != item.get('path')]
                        # remove checkbox
                        if cb in checks:
                            checks.remove(cb)
                        row_widget.setParent(None)
                    except Exception:
                        pass
                delete_btn.clicked.connect(on_delete)
                row.addWidget(thumb)
                row.addWidget(cb)
                row.addWidget(preview_btn)
                row.addWidget(delete_btn)
                row.addWidget(comm)
                rows.append((cb, comm))
                container_layout.addWidget(row_widget)
            note_label = QLabel("Comment:")
            v.addWidget(note_label)
            note_edit = QPlainTextEdit()
            try:
                note_edit.setFixedHeight(60)
            except Exception:
                pass
            v.addWidget(note_edit)
            btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
            v.addWidget(btns)
            def on_select_all(t):
                for c in checks:
                    c.setChecked(t)
            select_all.toggled.connect(on_select_all)
            def on_ok():
                note = note_edit.toPlainText().strip()
                if note:
                    self._insert_header_note(note)
                for c, comm in rows:
                    if c.isChecked():
                        d = c.item_data
                        note_to_use = comm.text().strip()
                        self._insert_image_with_note(d.get('path'), note_to_use, d.get('screen'), d.get('time'))
                dlg.accept()
            btns.accepted.connect(on_ok)
            btns.rejected.connect(dlg.reject)
            dlg.exec()
        except Exception:
            pass

    def _capture_tick(self):
        try:
            choice = self.screen_combo.currentText() if hasattr(self, 'screen_combo') else "All Screens"
            now = QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss")
            if choice == "All Screens":
                for s in QApplication.screens():
                    pix = s.grabWindow(0)
                    os.makedirs(self.images_dir, exist_ok=True)
                    path = os.path.join(self.images_dir, f"screencap_{uuid.uuid4().hex}.png")
                    pix.save(path, 'PNG')
                    self.captured_images.append({"path": path, "screen": s.name(), "time": now})
            else:
                target = None
                for s in QApplication.screens():
                    if s.name() == choice:
                        target = s
                        break
                if target:
                    pix = target.grabWindow(0)
                    os.makedirs(self.images_dir, exist_ok=True)
                    path = os.path.join(self.images_dir, f"screencap_{uuid.uuid4().hex}.png")
                    pix.save(path, 'PNG')
                    self.captured_images.append({"path": path, "screen": target.name(), "time": now})
        except Exception:
            pass

    def _capture_click(self):
        try:
            choice = self.screen_combo.currentText() if hasattr(self, 'screen_combo') else "All Screens"
            now = QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss")
            if choice == "All Screens":
                for s in QApplication.screens():
                    pix = s.grabWindow(0)
                    os.makedirs(self.images_dir, exist_ok=True)
                    path = os.path.join(self.images_dir, f"screencap_{uuid.uuid4().hex}.png")
                    pix.save(path, 'PNG')
                    self.captured_images.append({"path": path, "screen": s.name(), "time": now})
            else:
                target = None
                for s in QApplication.screens():
                    if s.name() == choice:
                        target = s
                        break
                if target:
                    pix = target.grabWindow(0)
                    os.makedirs(self.images_dir, exist_ok=True)
                    path = os.path.join(self.images_dir, f"screencap_{uuid.uuid4().hex}.png")
                    pix.save(path, 'PNG')
                    self.captured_images.append({"path": path, "screen": target.name(), "time": now})
            try:
                self.status_bar.showMessage("Captured", 1000)
            except Exception:
                pass
        except Exception:
            pass
        
    def toggle_always_on_top(self, checked):
        if checked:
            self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowStaysOnTopHint)
        else:
            self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowStaysOnTopHint)
        self.show()
    
    def toggle_fullscreen(self):
        if self.isFullScreen():
            self.showNormal()
        else:
            self.showFullScreen()
        
    def toggle_word_wrap(self, checked):
        if checked:
            self.current_tab().setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        else:
            self.current_tab().setLineWrapMode(QTextEdit.LineWrapMode.NoWrap)
    
    def toggle_line_numbers(self, checked):
        editor = self.current_tab()
        if hasattr(editor, "setLineNumbersVisible"):
            editor.setLineNumbersVisible(checked)
    
    def set_language(self, language):
        editor = self.current_tab()
        if hasattr(editor, "setLanguage"):
            editor.setLanguage(language)
        self._close_dev_view_if_present()
    
    def _close_dev_view_if_present(self):
        idx = self.tab_widget.currentIndex()
        w = self.tab_widget.widget(idx)
        if hasattr(w, 'text_editor'):
            title = self.tab_widget.tabText(idx)
            editor = w.text_editor
            proc = getattr(w, 'process', None)
            try:
                if proc and proc.state() != QProcess.ProcessState.NotRunning:
                    proc.kill()
                    proc.waitForFinished(1000)
            except Exception:
                pass
            self.tab_widget.removeTab(idx)
            self.tab_widget.insertTab(idx, editor, title)
            self.tab_widget.setCurrentIndex(idx)
    
    def open_dev_view(self, language):
        idx = self.tab_widget.currentIndex()
        title = self.tab_widget.tabText(idx)
        editor = self.current_tab()
        parent_widget = self.tab_widget.widget(idx)
        if hasattr(parent_widget, 'text_editor'):
            container_old = parent_widget
            proc_old = getattr(container_old, 'process', None)
            if proc_old and proc_old.state() != QProcess.ProcessState.NotRunning:
                proc_old.kill()
                proc_old.waitForFinished(2000)
            if hasattr(editor, 'setLanguage'):
                editor.setLanguage(language)
            splitter = getattr(container_old, 'splitter', None)
            term_widget = getattr(container_old, 'term_widget', None)
            if splitter and term_widget:
                splitter.insertWidget(0, editor)
                splitter.insertWidget(1, term_widget)
                splitter.setStretchFactor(0, 4)
                splitter.setStretchFactor(1, 1)
                h = max(400, self.height())
                QTimer.singleShot(0, lambda: splitter.setSizes([int(h*0.8), int(h*0.2)]))
            proc = getattr(container_old, 'process', None)
            if language == "Python" and proc:
                proc.start('python', ['-i', '-u', '-q'])
            elif language == "R" and proc:
                candidates = ['R', 'Rterm.exe', 'R.exe']
                exe = None
                for c in candidates:
                    if shutil.which(c):
                        exe = c
                        break
                if exe:
                    args = ['--no-save', '--quiet'] if 'Rterm' in exe else ['--vanilla', '--quiet']
                    proc.start(exe, args)
                else:
                    QMessageBox.information(self, 'R', 'R interpreter not found on PATH')
            return
        container = QWidget()
        layout = QVBoxLayout(container)
        splitter = QSplitter(Qt.Orientation.Vertical)
        splitter.addWidget(editor)
        term_widget = QWidget()
        term_layout = QVBoxLayout(term_widget)
        out_view = QPlainTextEdit()
        out_view.setReadOnly(True)
        out_view.setMinimumHeight(120)
        in_line = QLineEdit()
        in_line.setMinimumHeight(28)
        proc = QProcess(self)
        proc.setProcessChannelMode(QProcess.ProcessChannelMode.MergedChannels)
        def append_output():
            data = proc.readAllStandardOutput().data().decode('utf-8', errors='replace')
            out_view.appendPlainText(data)
        proc.readyReadStandardOutput.connect(append_output)
        in_line.returnPressed.connect(lambda: proc.write((in_line.text() + '\n').encode('utf-8')))
        term_layout.addWidget(out_view)
        term_layout.addWidget(in_line)
        splitter.addWidget(term_widget)
        splitter.insertWidget(0, editor)
        try:
            editor.setMinimumHeight(300)
        except Exception:
            pass
        splitter.setStretchFactor(0, 4)
        splitter.setStretchFactor(1, 1)
        h = max(400, self.height())
        QTimer.singleShot(0, lambda: splitter.setSizes([int(h*0.8), int(h*0.2)]))
        layout.addWidget(splitter)
        container.text_editor = editor
        container.process = proc
        container.output = out_view
        container.input = in_line
        container.splitter = splitter
        container.term_widget = term_widget
        self.tab_widget.removeTab(idx)
        self.tab_widget.insertTab(idx, container, title)
        self.tab_widget.setCurrentIndex(idx)
        if language == "Python":
            if proc.state() == QProcess.ProcessState.NotRunning:
                proc.start('python', ['-i', '-u', '-q'])
        elif language == "R":
            candidates = ['R', 'Rterm.exe', 'R.exe']
            exe = None
            for c in candidates:
                if shutil.which(c):
                    exe = c
                    break
            if exe:
                args = ['--no-save', '--quiet'] if 'Rterm' in exe else ['--vanilla', '--quiet']
                proc.start(exe, args)
            else:
                QMessageBox.information(self, 'R', 'R interpreter not found on PATH')
    
    def zoom_in(self):
        self.zoom_level += 10
        self.current_tab().zoomIn(1)
        self.update_status()
    
    def zoom_out(self):
        if self.zoom_level > 20:
            self.zoom_level -= 10
            self.current_tab().zoomOut(1)
            self.update_status()
    
    def reset_zoom(self):
        # Reset to 100%
        while self.zoom_level > 100:
            self.current_tab().zoomOut(1)
            self.zoom_level -= 10
        while self.zoom_level < 100:
            self.current_tab().zoomIn(1)
            self.zoom_level += 10
        self.update_status()
    
    def apply_theme(self, theme_name):
        theme = self.themes.get(theme_name, self.themes["Light"])
        palette = QPalette()
        palette.setColor(QPalette.ColorRole.Base, QColor(theme["bg"]))
        palette.setColor(QPalette.ColorRole.Text, QColor(theme["fg"]))
        self.current_tab().setPalette(palette)
        self.status_bar.showMessage(f"Applied {theme_name} theme", 2000)
    
    def show_word_count(self):
        text = self.current_tab().toPlainText()
        words = len(text.split())
        QMessageBox.information(self, "Word Count", f"Total words: {words}")
    
    def show_char_count(self):
        text = self.current_tab().toPlainText()
        chars = len(text)
        chars_no_space = len(text.replace(" ", "").replace("\n", ""))
        QMessageBox.information(self, "Character Count", 
                              f"Total characters: {chars}\nWithout spaces: {chars_no_space}")
    
    def convert_to_upper(self):
        cursor = self.current_tab().textCursor()
        if cursor.hasSelection():
            text = cursor.selectedText()
            cursor.insertText(text.upper())
        else:
            text = self.current_tab().toPlainText()
            self.current_tab().setPlainText(text.upper())
    
    def convert_to_lower(self):
        cursor = self.current_tab().textCursor()
        if cursor.hasSelection():
            text = cursor.selectedText()
            cursor.insertText(text.lower())
        else:
            text = self.current_tab().toPlainText()
            self.current_tab().setPlainText(text.lower())
    
    def convert_to_title(self):
        cursor = self.current_tab().textCursor()
        if cursor.hasSelection():
            text = cursor.selectedText()
            cursor.insertText(text.title())
        else:
            text = self.current_tab().toPlainText()
            self.current_tab().setPlainText(text.title())
    
    def remove_duplicate_lines(self):
        text = self.current_tab().toPlainText()
        lines = text.split('\n')
        unique_lines = []
        seen = set()
        for line in lines:
            if line not in seen:
                unique_lines.append(line)
                seen.add(line)
        self.current_tab().setPlainText('\n'.join(unique_lines))
        self.status_bar.showMessage(f"Removed {len(lines) - len(unique_lines)} duplicate lines", 3000)
    
    def sort_lines(self):
        text = self.current_tab().toPlainText()
        lines = text.split('\n')
        sorted_lines = sorted(lines)
        self.current_tab().setPlainText('\n'.join(sorted_lines))
        self.status_bar.showMessage("Lines sorted alphabetically", 2000)

    def correct_selection(self):
        editor = self.current_tab()
        if not hasattr(editor, 'highlighter'):
            return
        lang = getattr(editor.highlighter, 'language', 'Plain Text')
        cursor = editor.textCursor()
        if cursor.hasSelection():
            raw = cursor.selectedText()
        else:
            raw = editor.toPlainText()
        s = raw.replace('\u2029', '\n')
        corrected = s
        if cursor.hasSelection():
            cursor.insertText(corrected)
        else:
            editor.setPlainText(corrected)
        self.status_bar.showMessage('Selection corrected', 2000)

    

    def get_current_language(self):
        editor = self.current_tab()
        return getattr(editor.highlighter, 'language', 'Text') if hasattr(editor, 'highlighter') else 'Text'

    def show_run_output(self, text, title='Run Output'):
        dlg = QDialog(self)
        dlg.setWindowTitle(title)
        v = QVBoxLayout(dlg)
        out = QPlainTextEdit()
        out.setReadOnly(True)
        out.setPlainText(text)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        btns.rejected.connect(dlg.reject)
        v.addWidget(out)
        v.addWidget(btns)
        dlg.resize(800, 500)
        dlg.exec()

    def run_current(self):
        editor = self.current_tab()
        lang = self.get_current_language()
        text = editor.toPlainText()
        if lang == 'Python':
            try:
                with open(self.python_temp_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                proc = subprocess.run(['python', self.python_temp_path], capture_output=True, text=True)
                self.show_run_output(proc.stdout + proc.stderr, 'Python Run')
            except Exception as e:
                QMessageBox.warning(self, 'Run', str(e))
        elif lang == 'R':
            try:
                with open(self.r_temp_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                exe = shutil.which('Rscript') or shutil.which('Rscript.exe')
                if exe:
                    args = [exe, self.r_temp_path]
                else:
                    exe = shutil.which('R') or shutil.which('R.exe') or shutil.which('Rterm.exe')
                    if not exe:
                        QMessageBox.information(self, 'Run', 'Rscript not found')
                        return
                    args = [exe, '--vanilla', '--quiet', '-f', self.r_temp_path]
                proc = subprocess.run(args, capture_output=True, text=True)
                self.show_run_output(proc.stdout + proc.stderr, 'R Run')
            except Exception as e:
                QMessageBox.warning(self, 'Run', str(e))
        else:
            QMessageBox.information(self, 'Run', 'Select Python or R language')

    def run_selection(self):
        editor = self.current_tab()
        lang = self.get_current_language()
        cursor = editor.textCursor()
        if cursor.hasSelection():
            text = cursor.selectedText().replace('\u2029', '\n')
        else:
            text = editor.toPlainText()
        if lang == 'Python':
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.py', mode='w', encoding='utf-8') as tf:
                    tf.write(text)
                    path = tf.name
                proc = subprocess.run(['python', path], capture_output=True, text=True)
                try:
                    os.unlink(path)
                except Exception:
                    pass
                self.show_run_output(proc.stdout + proc.stderr, 'Python Selection')
            except Exception as e:
                QMessageBox.warning(self, 'Run', str(e))
        elif lang == 'R':
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.R', mode='w', encoding='utf-8') as tf:
                    tf.write(text)
                    path = tf.name
                exe = shutil.which('Rscript') or shutil.which('Rscript.exe')
                if exe:
                    args = [exe, path]
                else:
                    exe = shutil.which('R') or shutil.which('R.exe') or shutil.which('Rterm.exe')
                    if not exe:
                        QMessageBox.information(self, 'Run', 'Rscript not found')
                        return
                    args = [exe, '--vanilla', '--quiet', '-f', path]
                proc = subprocess.run(args, capture_output=True, text=True)
                try:
                    os.unlink(path)
                except Exception:
                    pass
                self.show_run_output(proc.stdout + proc.stderr, 'R Selection')
            except Exception as e:
                QMessageBox.warning(self, 'Run', str(e))
        else:
            QMessageBox.information(self, 'Run', 'Select Python or R language')

    def export_text(self):
        editor = self.current_tab()
        path, _ = QFileDialog.getSaveFileName(self, 'Export Text', '', 'Text Files (*.txt);;All Files (*.*)')
        if not path:
            return
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(editor.toPlainText())
            self.status_bar.showMessage('Exported to text', 3000)
        except Exception as e:
            QMessageBox.critical(self, 'Export', str(e))

    def export_markdown(self):
        editor = self.current_tab()
        path, _ = QFileDialog.getSaveFileName(self, 'Export Markdown', '', 'Markdown Files (*.md);;All Files (*.*)')
        if not path:
            return
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(editor.toPlainText())
            self.status_bar.showMessage('Exported to markdown', 3000)
        except Exception as e:
            QMessageBox.critical(self, 'Export', str(e))

    def export_html(self):
        path, _ = QFileDialog.getSaveFileName(self, 'Export HTML', '', 'HTML Files (*.html);;All Files (*.*)')
        if not path:
            return
        try:
            html = self.current_tab().document().toHtml()
            base = os.path.splitext(os.path.basename(path))[0]
            assets = os.path.join(os.path.dirname(path), base + '_assets')
            os.makedirs(assets, exist_ok=True)
            html2 = self._rewrite_html_img_src_for_export(html, assets)
            with open(path, 'w', encoding='utf-8') as f:
                f.write(html2)
            self.status_bar.showMessage('Exported to HTML', 3000)
        except Exception as e:
            QMessageBox.critical(self, 'Export', str(e))

    def export_pdf(self):
        path, _ = QFileDialog.getSaveFileName(self, 'Export PDF', '', 'PDF Files (*.pdf);;All Files (*.*)')
        if not path:
            return
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(path)
            self.current_tab().document().print(printer)
            self.status_bar.showMessage('Exported to PDF', 3000)
        except Exception as e:
            QMessageBox.critical(self, 'Export', str(e))

    def _rewrite_html_img_src_for_export(self, html, assets_dir):
        try:
            import re
            from urllib.parse import urlparse, unquote
            os.makedirs(assets_dir, exist_ok=True)
            def repl(m):
                src = m.group(1)
                if src.startswith('data:'):
                    return m.group(0)
                p = src
                try:
                    u = urlparse(src)
                    if u.scheme in ('file', ''):
                        p = unquote(u.path) if u.scheme == 'file' else src
                except Exception:
                    p = src
                if not os.path.isabs(p):
                    p = src
                if not os.path.exists(p):
                    return m.group(0)
                ext = os.path.splitext(p)[1] or '.png'
                dest = os.path.join(assets_dir, f"img_{uuid.uuid4().hex}{ext}")
                try:
                    shutil.copy2(p, dest)
                    rel = os.path.basename(dest)
                    return f'src="{rel}"'
                except Exception:
                    return m.group(0)
            return re.sub(r'src="([^"]+)"', repl, html)
        except Exception:
            return html

    def create_sample_files(self):
        os.makedirs(self.session_dir, exist_ok=True)
        samples = []
        csv_path = os.path.join(self.session_dir, 'sample.csv')
        with open(csv_path, 'w', encoding='utf-8') as f:
            f.write('id,name,score\n1,Alice,91\n2,Bob,88\n3,Charlie,95\n')
        samples.append(csv_path)
        tsv_path = os.path.join(self.session_dir, 'sample.tsv')
        with open(tsv_path, 'w', encoding='utf-8') as f:
            f.write('id\tcity\tzip\n1\tNY\t10001\n2\tSF\t94105\n')
        samples.append(tsv_path)
        txt_path = os.path.join(self.session_dir, 'sample.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('This is a plain text sample file.')
        samples.append(txt_path)
        md_path = os.path.join(self.session_dir, 'sample.md')
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('# Sample Markdown\n\n| id | name |\n|---|---|\n| 1 | Alice |\n| 2 | Bob |')
        samples.append(md_path)
        json_path = os.path.join(self.session_dir, 'sample.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump({"users": [{"id": 1, "name": "Alice"}, {"id": 2, "name": "Bob"}]}, f, indent=2)
        samples.append(json_path)
        jsonl_path = os.path.join(self.session_dir, 'sample.jsonl')
        with open(jsonl_path, 'w', encoding='utf-8') as f:
            f.write('{"id":1,"val":10}\n{"id":2,"val":20}\n')
        samples.append(jsonl_path)
        py_path = os.path.join(self.session_dir, 'sample.py')
        with open(py_path, 'w', encoding='utf-8') as f:
            f.write('print("Hello from Python sample")\n')
        samples.append(py_path)
        r_path = os.path.join(self.session_dir, 'sample.R')
        with open(r_path, 'w', encoding='utf-8') as f:
            f.write('print("Hello from R sample")\n')
        samples.append(r_path)
        pdf_path = os.path.join(self.session_dir, 'sample.pdf')
        try:
            minimal_pdf = b"%PDF-1.4\n1 0 obj<<>>endobj\n2 0 obj<<>>endobj\n3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 300 200]/Contents 4 0 R>>endobj\n4 0 obj<</Length 55>>stream\nBT /F1 24 Tf 50 150 Td (Sample PDF Page) Tj ET\nendstream\nendobj\n5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n6 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n7 0 obj<</Type/Catalog/Pages 6 0 R>>endobj\nxref\n0 8\n0000000000 65535 f \n0000000010 00000 n \n0000000045 00000 n \n0000000068 00000 n \n0000000174 00000 n \n0000000321 00000 n \n0000000386 00000 n \n0000000447 00000 n \ntrailer<</Size 8/Root 7 0 R>>\nstartxref\n505\n%%EOF"
            with open(pdf_path, 'wb') as f:
                f.write(minimal_pdf)
            samples.append(pdf_path)
        except Exception:
            pass
        if pl and hasattr(pl.DataFrame({"a":[1]}), 'write_excel'):
            xlsx_path = os.path.join(self.session_dir, 'sample.xlsx')
            try:
                pl.DataFrame({"id":[1,2],"name":["Alice","Bob"]}).write_excel(xlsx_path)
                samples.append(xlsx_path)
            except Exception:
                pass
        for p in samples:
            try:
                self.open_specific_file(p)
            except Exception:
                pass
        QMessageBox.information(self, 'Samples', 'Sample files created and opened')

    def open_specific_file(self, path):
        try:
            self.current_file = None
            self.is_modified = False
            self.update_title()
            self.tab_widget.setCurrentWidget(self.new_tab())
            self.open_file_path(path)
        except Exception:
            pass

    def open_file_path(self, filename):
        try:
            editor = self.current_tab()
            lower = filename.lower()
            if lower.endswith(('.csv', '.tsv', '.txt')) and duckdb:
                tmp = os.path.join(self.session_dir, f"duckdb_{os.path.basename(filename)}.txt")
                os.makedirs(self.session_dir, exist_ok=True)
                try:
                    con = duckdb.connect()
                    con.execute(f"COPY (SELECT * FROM read_csv_auto('{filename}')) TO '{tmp}' (DELIMITER '|', HEADER TRUE)")
                    with open(tmp, 'r', encoding='utf-8', errors='replace') as f:
                        editor.setPlainText(f.read())
                except Exception:
                    import csv
                    with open(filename, 'r', encoding='utf-8', errors='replace', newline='') as fin, open(tmp, 'w', encoding='utf-8', newline='') as fout:
                        sample = fin.read(4096)
                        fin.seek(0)
                        try:
                            dialect = csv.Sniffer().sniff(sample)
                        except Exception:
                            dialect = csv.excel
                        reader = csv.reader(fin, dialect)
                        w = csv.writer(fout, delimiter='|')
                        for row in reader:
                            w.writerow(row)
                    with open(tmp, 'r', encoding='utf-8', errors='replace') as f:
                        editor.setPlainText(f.read())
            elif lower.endswith(('.xlsx', '.xls')) and pl:
                df = pl.read_excel(filename)
                tmp = os.path.join(self.session_dir, f"polars_{os.path.basename(filename)}.txt")
                os.makedirs(self.session_dir, exist_ok=True)
                df.write_csv(tmp, separator='|')
                with open(tmp, 'r', encoding='utf-8', errors='replace') as f:
                    editor.setPlainText(f.read())
            elif lower.endswith(('.json', '.ndjson', '.jsonl')):
                with open(filename, 'r', encoding='utf-8', errors='replace') as f:
                    raw = f.read()
                try:
                    obj = json.loads(raw)
                    editor.setPlainText(json.dumps(obj, indent=2, ensure_ascii=False))
                except Exception:
                    lines = []
                    for line in raw.splitlines():
                        try:
                            lines.append(json.dumps(json.loads(line), indent=2, ensure_ascii=False))
                        except Exception:
                            lines.append(line)
                    editor.setPlainText("\n".join(lines))
            elif lower.endswith(('.pdf')) and PdfReader:
                reader = PdfReader(filename)
                pages = []
                for p in reader.pages:
                    pages.append(p.extract_text() or '')
                editor.setPlainText("\n\n".join(pages))
            else:
                with open(filename, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                if lower.endswith(('.html', '.htm')):
                    editor.setHtml(content)
                else:
                    editor.setPlainText(content)
            if lower.endswith('.py') and hasattr(editor, 'setLanguage'):
                editor.setLanguage('Python')
            elif lower.endswith('.r') and hasattr(editor, 'setLanguage'):
                editor.setLanguage('R')
            elif lower.endswith('.md') and hasattr(editor, 'setLanguage'):
                editor.setLanguage('Text')
            self.tab_widget.setTabText(self.tab_widget.currentIndex(), os.path.basename(filename))
        except Exception:
            pass


def main():
    app = QApplication(sys.argv)
    app.setApplicationName("Ultra Advanced Notepad++")
    window = AdvancedNotepad()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
